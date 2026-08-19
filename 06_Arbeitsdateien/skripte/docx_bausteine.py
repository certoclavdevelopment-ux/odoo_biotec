#!/usr/bin/env python3
"""Gemeinsame Word-Bausteine für die Dokumentgeneratoren dieses Projekts.

Layout, Farben und Kopf-/Fußzeile der CertoClav-Dokumente an einer Stelle, damit
Protokolle und Kundendokumente gleich aussehen.
"""
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x4E, 0x79)      # dunkelblau, Deck-Akzent
GREY = RGBColor(0x7F, 0x7F, 0x7F)
OPEN_MARK = RGBColor(0xB0, 0x4A, 0x00)   # Hinweisfarbe für offene Punkte

FOOTER_1 = ("CertoClav Sterilizer GmbH  ·  Peintner Straße 10, 4060 Leonding, Österreich  ·  "
            "Geschäftsführer: Michael Simon (geb. Dirix), MSc.")
FOOTER_2 = ("Firmenbuch: Landesgericht Linz, FN 122912d  ·  UID ATU22821702  ·  "
            "Tel. +43 732 674 278  ·  support@certoclav.com  ·  www.certoclav.com")

OFFEN = "__OFFEN__"   # Marker: Frage nicht beantwortet


# --------------------------------------------------------------------------- Hilfsfunktionen
def set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.08


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def bottom_border(par, color="BFBFBF", size=6):
    pPr = par._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    pPr.append(borders)


def answer_lines(doc, count=2):
    """Beschreibbare Zeilen mit dezenter Linie darunter."""
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.4)
        bottom_border(p)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.add_run("– " + item)


def offen_hinweis(doc, text="Im Termin nicht beantwortet – offen."):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.space_before = Pt(1)
    r = p.add_run("– " + text)
    r.font.color.rgb = OPEN_MARK
    r.font.italic = True


def heading(doc, text, kicker=None):
    if kicker:
        k = doc.add_paragraph()
        k.paragraph_format.space_before = Pt(14)
        k.paragraph_format.space_after = Pt(0)
        r = k.add_run(kicker.upper())
        r.font.size = Pt(7.5)
        r.font.bold = True
        r.font.color.rgb = GREY
        r.font.name = "Calibri"
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(2) if kicker else Pt(14)
    h.paragraph_format.space_after = Pt(6)
    bottom_border(h, color="1F4E79", size=8)
    r = h.add_run(text)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    return h


def question(doc, text, hint=None, lines=2, antwort=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.bold = True
    if hint:
        h = p.add_run("  " + hint)
        h.font.italic = True
        h.font.color.rgb = GREY
        h.font.size = Pt(9)
    if antwort == OFFEN:
        offen_hinweis(doc)
    elif antwort:
        bullets(doc, antwort if isinstance(antwort, list) else [antwort])
    else:
        answer_lines(doc, lines)


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY


def table(doc, cols, widths, rows, header_fill="1F4E79", zeilen_fuellung=None):
    """Tabelle mit Kopfzeile.

    zeilen_fuellung: optional eine Funktion(zeile) -> Hex-Farbe oder None.
    Gibt sie eine Farbe zurück, wird die ganze Zeile damit hinterlegt.
    """
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (name, w) in enumerate(zip(cols, widths)):
        c = t.rows[0].cells[i]
        c.width = Cm(w)
        shade(c, header_fill)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(name)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        fuellung = zeilen_fuellung(row) if zeilen_fuellung else None
        for i, val in enumerate(row):
            cells[i].width = Cm(widths[i])
            if fuellung:
                shade(cells[i], fuellung)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            for k, teil in enumerate(str(val).split("**")):
                if not teil:
                    continue
                r = p.add_run(teil)
                r.font.size = Pt(9.5)
                r.font.bold = k % 2 == 1
    return t


def page_field(par):
    par.add_run("Seite ")
    for instr in ("PAGE", "NUMPAGES"):
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), instr)
        run = OxmlElement("w:r")
        txt = OxmlElement("w:t")
        txt.text = "1"
        run.append(txt)
        fld.append(run)
        par._p.append(fld)
        if instr == "PAGE":
            par.add_run(" von ")


def seite_einrichten(doc, kopf_rechts):
    """A4, Ränder, Kopf- und Fußzeile im CertoClav-Layout."""
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = sec.right_margin = Cm(2.0)

    hp = sec.header.paragraphs[0]
    hp.text = ""
    r = hp.add_run("CERTOCLAV CONSULTING")
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    t = hp.add_run("\t" + kopf_rechts)
    t.font.size = Pt(8)
    t.font.color.rgb = GREY
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bottom_border(hp)

    fp = sec.footer.paragraphs[0]
    fp.text = ""
    for i, zeile in enumerate((FOOTER_1, FOOTER_2)):
        ziel = fp if i == 0 else sec.footer.add_paragraph()
        ziel.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ziel.paragraph_format.space_after = Pt(0)
        r = ziel.add_run(zeile)
        r.font.size = Pt(6.5)
        r.font.color.rgb = GREY

    sp = sec.footer.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_before = Pt(2)
    page_field(sp)
    for run in sp.runs:
        run.font.size = Pt(7)
        run.font.color.rgb = GREY

    return sec


def titel(doc, haupt, unter=None):
    t = doc.add_paragraph()
    t.paragraph_format.space_after = Pt(0)
    r = t.add_run(haupt)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    if unter:
        s = doc.add_paragraph()
        s.paragraph_format.space_after = Pt(14)
        r = s.add_run(unter)
        r.font.size = Pt(11)
        r.font.color.rgb = GREY
