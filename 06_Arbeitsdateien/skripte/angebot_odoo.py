#!/usr/bin/env python3
"""Angebot Odoo-Einführung für die biotec GmbH.

    python3 angebot_odoo.py <ziel.docx>

Struktur nach der Work-Breakdown-Structure 1–8 aus
01_Projektsteuerung/todo-angebot-begleitdokument.md, damit biotec den
aktivierungsfähigen Anteil sauber vom Aufwand trennen kann.
"""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, __file__.rsplit("/", 1)[0])
from docx_bausteine import (ACCENT, GREY, seite_einrichten, titel, heading,
                            table, bullets, note, set_base_style)

DOK = "ANG-2026-001"
STAND = "Version 2, Stand 01.09.2026"
STUNDENSATZ = 137.50
STD_JE_TAG = 8
PT_SATZ = STUNDENSATZ * STD_JE_TAG          # 1.100,00 € je Personentag

GELB = "FFF2A8"
BLAU = "DCE6F1"


def eur(x):
    return f"{x:,.2f} €".replace(",", "§").replace(".", ",").replace("§", ".")


def pt(x):
    return f"{x:,.1f}".replace(".", ",")


def absatz(doc, text, groesse=10, kursiv=False, abstand=6, fett=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(abstand)
    for i, teil in enumerate(text.split("**")):
        if not teil:
            continue
        r = p.add_run(teil)
        r.font.size = Pt(groesse)
        r.font.italic = kursiv
        r.font.bold = fett or i % 2 == 1
    return p


# ---------------------------------------------------------------- Leistungen
# (Nr., Inhalt, PT, Einordnung, Phase)
AKT = "aktivierungsbezogen"
AUF = "Aufwand"

WORKSTREAMS = [
    ("1 · Analyse, Konzeption und Projektleitung", [
        ("1.1", "Prozessaufnahme Labor: Auftrag, Probenahme, Bebrütung, Auszählung, "
                "Gutachtenübergabe, Rechnung – an allen drei Standorten", 5.0, AUF, 1),
        ("1.2", "Prozessaufnahme Schulungsgeschäft und kaufmännische Abläufe. Für das "
                "Schulungsgeschäft gibt es heute kein System und keine Dokumentation", 4.0, AUF, 1),
        ("1.3", "Grobkonzept: Abbildungsentscheidungen je Bereich, abgestimmt und freigegeben",
         3.0, AUF, 1),
        ("1.4", "Detailspezifikation je Modul – prüffähige Anforderungsbeschreibung und "
                "Grundlage der Abnahme", 4.0, AKT, 1),
        ("1.5", "Rollen- und Rechtekonzept, Grundlage für Lizenzzahl und Zugriffe", 2.0, AUF, 1),
        ("1.6", "Auftaktveranstaltung: Vorbereitung und Durchführung", 1.0, AUF, 1),
        ("1.7", "Projektleitung, Abstimmungen, Statusberichte über die Laufzeit", 20.0, AUF, 1),
        ("1.8", "Projektabschluss, Übergabeverzeichnis der Arbeitsergebnisse, "
                "Abnahmeprotokolle", 2.0, AUF, 1),
    ]),
    ("2 · Konfiguration der Odoo-Module", [
        ("2.1", "Grundeinrichtung: Unternehmen, drei Standorte, Nummernkreise lückenlos, "
                "Benutzer, Sprachen, Fremdwährung CHF mit Kursquelle und Stichtagsbewertung",
         4.5, AKT, 1),
        ("2.2", "Kontakte: Kundenkategorien, Ansprechpartner, Werke und Gebäude als "
                "untergeordnete Kontakte", 3.0, AKT, 1),
        ("2.3", "Verkauf: Preislisten für Prüfleistungen, Kurse und Inhouse-Schulungen, "
                "Angebots- und Auftragsarten", 4.0, AKT, 1),
        ("2.4", "Rahmenverträge und wiederkehrende Abrechnung: Preisstaffeln, Jahresturnus, "
                "Sammelangebote je Werk", 3.0, AKT, 1),
        ("2.5", "Leistungs- und Artikelstamm: Prüfleistungen, Kursarten, Verbrauchsmaterial, "
                "Kategorien mit Erlöskonten", 2.5, AKT, 1),
        ("2.6", "Einkauf und Lager: Lagerorte, Verbrauchsmaterial, Nachbestellregeln für "
                "Gelatinefilter und Nährmedien", 3.0, AKT, 1),
        ("2.7", "**Buchhaltung Grundeinrichtung:** SKR04 mit Sachkontenlänge 4, "
                "Personenkonten, Journale, Steuerschlüssel 19 und 7 Prozent, § 13b als "
                "Leistungsempfänger für externe Dozenten, Drittland Schweiz, GoBD-Sperrdaten",
         6.0, AKT, 1),
        ("2.8", "Zahlungsverkehr: Kontoauszugsimport CAMT oder MT940, SEPA-Überweisungsdateien, "
                "Abgleichregeln, Skonto, Teilzahlungen", 4.0, AKT, 1),
        ("2.9", "Mahnwesen: drei Stufen, Texte, Fristen, Gebühren, Sperrlogik, Layout",
         2.5, AKT, 1),
        ("2.10", "Umsatzsteuervoranmeldung: Auswertung, Abstimmung, Übergabeweg. Die "
                 "Übermittlung an ELSTER erfolgt außerhalb von Odoo", 1.5, AKT, 1),
        ("2.11", "Anzahlungen und Vorkasse einschließlich Kursgebühren, Storno- und "
                 "Umbuchungsregeln", 2.0, AKT, 1),
        ("2.12", "**E-Rechnung:** Ausgang als Factur-X und XRechnung, Eingangsverarbeitung, "
                 "Anbindung an das Peppol-Netz", 4.0, AKT, 1),
        ("2.13", "Eingangsrechnungsverarbeitung mit Freigabeweg", 2.0, AKT, 1),
        ("2.14", "Analytische Buchhaltung: Standort, Geschäftsbereich und Projekt als drei "
                 "getrennte Dimensionen", 3.5, AKT, 1),
        ("2.15", "**RLT-Anlagenstamm und Turnussteuerung:** Gerätekategorien, Prüfintervalle "
                 "je Anlagentyp, automatisch erzeugte Wartungsanfragen mit Fälligkeit",
         5.0, AKT, 1),
        ("2.16", "Projekte und Zeiterfassung: Projektstruktur je Auftrag, analytische "
                 "Verteilung, Genehmigungsweg, Zeit zu Rechnung für Inhouse-Schulungen und "
                 "Analytik, Erfassung über das Mobilgerät", 4.0, AKT, 1),
        ("2.17", "Abwesenheiten und Urlaub – Grundlage für Einsatzplanung und Rückstellungen",
         1.5, AKT, 1),
        ("2.18", "Fremdleistungseinkauf mit Auftragsbezug und periodengerechter Abgrenzung",
         2.5, AKT, 1),
        ("2.19", "Spesen und Reisekosten einschließlich Weiterbelastung an den Kunden",
         2.0, AKT, 1),
        ("2.20", "Anlagenliste und Probenahmeauftrag als Druckausgabe aus Odoo – ersetzt die "
                 "heutige Zusatzsoftware", 2.0, AKT, 1),
        ("2.21", "**Anlagenbuchhaltung:** Anlagenklassen, Abschreibungsmethoden, "
                 "Anlagenspiegel, steuerliche und wirtschaftliche Nutzungsdauer getrennt",
         3.0, AKT, 1),
        ("2.22", "Rechte in Odoo umsetzen und Rechteprobe je Rolle protokollieren. Geteilte "
                 "Sammelkonten wie heute im Labor sind in einem Buchhaltungssystem nicht "
                 "haltbar", 3.0, AKT, 1),
        ("2.23", "Lösch- und Aufbewahrungskonzept, Beitrag zum Verzeichnis der "
                 "Verarbeitungstätigkeiten", 1.5, AUF, 1),
        ("2.24", "Dokumentenablage: Entscheidung zwischen OneDrive und Odoo Dokumente, "
                 "Verknüpfung der Gutachten am Auftrag, revisionssichere Ablage", 3.0, AKT, 1),
        ("2.25", "E-Mail-Betrieb: eigene Absenderdomäne, SPF, DKIM und DMARC, "
                 "Rückläuferbehandlung, Eingangsadressen, Vorlagen", 2.5, AKT, 1),
        ("2.26", "Test- und Abnahmeumgebung, Umgebungs- und Versionskonzept", 2.0, AKT, 1),
        ("2.27", "**Schulungsgeschäft:** Kursarten VDI 6022 und VDI 2047, Seminartermine mit "
                 "Ort und Kapazität, Mindestteilnehmerzahl, Anmeldung, Bestätigungs- und "
                 "Erinnerungsmails, Dozenten und Honorare", 8.0, AKT, 2),
    ]),
    ("3 · Kundenspezifische Entwicklung", [
        ("3.1", "Belegformulare im biotec-Layout: Angebot, Auftragsbestätigung, Lieferschein, "
                "Rechnung, Gutschrift, Mahnung, Bestellung – sieben Belegarten", 9.0, AKT, 1),
        ("3.2", "**Erweiterungsmodul `biotec_rlt`:** VDI-Luftwerte, drei Kennungen je Anlage, "
                "Standortfelder, Prüfhistorie – eigenes Modul mit Quellcode", 8.0, AKT, 1),
        ("3.3", "Arbeitsvorrat und Fälligkeitsübersicht der Wiederholungsprüfungen, "
                "Terminbündelung je Werk, Sammelangebot statt Einzelangeboten", 6.0, AKT, 1),
        ("3.4", "Rückmeldekanal „Gutachten fertiggestellt" + chr(8220) + " aus dem Altsystem "
                "nach Odoo als Auslöser der Rechnungsfreigabe", 3.0, AKT, 1),
        ("3.5", "Auswertungen und Kennzahlen: Fälligkeiten, Auslastung Labor, Umsatz je "
                "Bereich und Standort, offene Posten", 4.0, AKT, 1),
        ("3.6", "Teilnahmebescheinigung und Zertifikatsvorlage mit Nummer, Gültigkeit und "
                "Sammeldruck", 3.0, AKT, 2),
    ]),
    ("4 · Schnittstellen", [
        ("4.1", "**Stammdatenbrücke Odoo zum Gutachten-Altsystem:** Analyse der "
                "Feldbedeutung, Anbindung, Zeichenkodierung, Erstabgleich gegen eine Kopie, "
                "Dienstbetrieb mit Protokoll, Neuanlagen samt Nummernvergabe, Anpassung der "
                "Odoo-Oberfläche, gemeinsamer Einstieg, Abnahme", 17.0, AKT, 1),
        ("4.2", "DATEV-Übergabe: Formatabstimmung, Testdatei, Abnahme mit dem Steuerberater",
         3.0, AKT, 1),
    ]),
    ("5 · Datenmigration", [
        ("5.1", "Kunden, Ansprechpartner, Werke und Gebäude – technische Übernahme und "
                "Neuaufbau der Struktur", 5.0, AKT, 1),
        ("5.2", "Anlagenstamm mit 938 RLT-Anlagen – technische Übernahme", 5.0, AKT, 1),
        ("5.3", "Lieferanten und Fremdlabore – Neuaufbau aus den Eingangsrechnungen der "
                "letzten zwölf Monate", 2.0, AKT, 1),
        ("5.4", "Leistungen, Artikel, Verbrauchsmaterial und Kurskatalog", 2.0, AKT, 1),
        ("5.5", "Kontenrahmen, offene Posten, Saldenvortrag zum Stichtag", 4.0, AKT, 1),
        ("5.6", "Einbuchung der Stichtagsinventur – gezählt wird von biotec", 1.5, AKT, 1),
        ("5.7", "Importmethodik mit externen Kennungen, Probeläufe, Abstimmung der Summen "
                "mit der Buchhaltung", 3.0, AKT, 1),
        ("5.8", "**Letztes Prüfdatum je Anlage ermitteln und einpflegen** – aus den "
                "versionierten Prüflisten und dem Gutachtenarchiv. Ohne dieses Datum erzeugt "
                "die Turnussteuerung am ersten Tag keine Fälligkeiten", 4.0, AKT, 1),
        ("5.9", "Inhaltliche Datenbereinigung: Kundenliste, 68 Herstellerschreibweisen, "
                "27 Wartungsfirmen, Testdatensätze, Zahlenformate. Entfällt ganz oder "
                "teilweise, wenn biotec die Bereinigung nach unserer Vorlage selbst "
                "übernimmt – siehe M12", 8.0, AUF, 1),
    ]),
    ("6 · Integration und Test", [
        ("6.1", "Funktionstests je Modul, rund 60 Testfälle mit dokumentiertem Ergebnis",
         12.0, AKT, 1),
        ("6.2", "Durchgängige Prozesstests: Auftrag bis Rechnung, Kurs von der Anmeldung bis "
                "zur Bescheinigung, Einkauf bis Lieferantenrechnung", 6.0, AKT, 1),
        ("6.3", "Betriebssimulation über einen Monat mit dem Fachbereich, mit echten Fällen",
         8.0, AUF, 1),
        ("6.4", "Fehlerbehebung und Abnahmeprotokolle je Arbeitspaket", 6.0, AKT, 1),
    ]),
    ("7 · Schulung und Befähigung", [
        ("7.1", "Anwenderleitfäden für sechs Rollen: Labor, Probennehmer, Verwaltung, "
                "Buchhaltung, Schulungsorganisation, Geschäftsführung", 6.0, AUF, 1),
        ("7.2", "Vorbereitung: Übungsmandant, Übungsdaten, Ablaufplan", 2.0, AUF, 1),
        ("7.3", "**Vor-Ort-Schulung in Gütersloh, drei Tage:** zwei Tage für die vier "
                "Kernnutzer, ein Tag für die weiteren Rollen", 3.0, AUF, 1),
        ("7.4", "Zusatzblock für Labor und Probennehmer", 2.0, AUF, 1),
        ("7.5", "Einweisung der Standorte München und Mittweida aus der Ferne", 1.5, AUF, 1),
        ("7.6", "An- und Abreise Leonding – Gütersloh", 2.0, AUF, 1),
        ("7.7", "Nachschulung vier Wochen nach dem Start, aus der Ferne", 2.0, AUF, 1),
    ]),
    ("8 · Umstellung und Betreuung", [
        ("8.1", "Umstellungsplan, Begleitung der Stichtagsinventur, abschließende "
                "Datenübernahme", 4.0, AUF, 1),
        ("8.2", "Begleitung am Umstellungstag und an den beiden Folgetagen", 2.0, AUF, 1),
        ("8.3", "Enge Betreuung über vier Wochen nach dem Start", 9.0, AUF, 1),
        ("8.4", "Begleitung des ersten Monatsabschlusses in Odoo", 3.0, AUF, 1),
        ("8.5", "GoBD-Verfahrensdokumentation und Migrationsdokumentation", 3.0, AKT, 1),
        ("8.6", "Betriebshandbuch, Sicherungskonzept und eine getestete Rücksicherung",
         2.0, AUF, 1),
    ]),
]

OPTIONEN = [
    ("O1", "Übernahme des Interessentenbestands mit 10.387 Adressen in das CRM, "
           "Dublettenabgleich gegen die Kunden, Prüfung der Rechtsgrundlage nach DSGVO", 6.0),
    ("O2a", "**Datenlieferung an LucaNet, wenn ein Lieferformat der Gruppe vorliegt:** "
            "Einrichtung, Testlieferung, Abnahme", 8.0),
    ("O2b", "Datenlieferung an LucaNet **ohne bestehendes Format:** zusätzlich Kontenmapping "
            "SKR04 auf den Gruppenkontenplan, Prüfsummen, Bilanzgleichung, Lieferprotokoll",
     14.0),
    ("O3a", "IFRS-Überleitung: Journal, Kontenzuordnung, Auswertungen", 8.0),
    ("O3b", "IFRS 16: Vertragsregister für Miet-, Leasing- und Fahrzeugverträge", 4.0),
    ("O3c", "IAS 38: Erfassung der Entwicklungsprojekte mit zurechenbaren Kosten", 3.0),
    ("O4", "Verknüpfung der vorhandenen Gutachten-PDF aus der OneDrive-Ablage am "
           "Anlagensatz, soweit eindeutig zuordenbar. Die LaTeX-Quellen des Altsystems "
           "werden nicht neu gesetzt", 6.0),
    ("O5a", "**POC/WIP – Buchhaltungsspezifikation** mit Steuerberater und "
            "Wirtschaftsprüfer: Methode, Konten, Verlustantizipation, Periodenabgrenzung",
     6.0),
    ("O5b", "**POC/WIP – Rechenkern:** Odoo-Modul mit Vertrags- und Periodenmodell, externer "
            "Rechenlauf mit Entwurfsbuchungen, Wiederholbarkeit, Storno, Bericht", 28.0),
    ("O6", "Online-Anmeldung für Seminare auf der Website mit Anbindung an Odoo", 4.0),
    ("O7", "**Umsetzung im Gutachtenprogramm durch CertoClav**, falls Westbomke EDV nicht "
           "beauftragt wird. Richtwert, Bezifferung erst nach Sichtung", 15.0),
    ("O8", "Prüfmittelüberwachung und Kalibrierung – Pflicht bei einer Akkreditierung nach "
           "DIN EN ISO/IEC 17025", 3.0),
    ("O9", "Einsatz- und Kapazitätsplanung des Prüferteams", 4.0),
    ("O10", "Nachweisregister der Schulungen mit Auffrischungserinnerung an die Teilnehmer",
     3.0),
    ("O11", "Chargen und Haltbarkeit für Nährmedien und Gelatinefilter", 2.0),
    ("O12", "Reklamationsbearbeitung", 1.5),
    ("O13", "**Anwendungsbetreuung ab dem Ende der Betreuungsphase:** 1,5 PT je Monat, "
            "zusätzlich 4 PT je Jahr für die Versionspflege der Eigenmodule und der Brücke",
     0.0),
]

MITWIRKUNG = [
    ["M1", "**Stichtagsinventur** des Verbrauchsmaterials unmittelbar vor der Umstellung – "
           "Gelatinefilter, Nährmedien, Röhrchen. Eine vollständige Lagerliste besteht heute "
           "nicht; der Bestand wird gezählt, nicht migriert", "biotec", "vor Umstellung"],
    ["M2", "**Finale Unterlagen zum Stichtag:** offene Posten Debitoren und Kreditoren, "
           "Saldenvortrag, letzter Kontostand", "Buchhaltung", "Stichtag"],
    ["M3", "**Summen- und Saldenliste als Excel oder CSV.** Die vorliegende Fassung besteht "
           "aus Bildschirmfotos und ist maschinell nicht auswertbar", "Buchhaltung",
     "Projektstart"],
    ["M4", "Fehlende Quelldateien des Altsystems, vor allem die Briefköpfe und Vorlagen aus "
           "der Gutachten-Software", "Westbomke EDV", "vor Layoutarbeit"],
    ["M5", "Prüfintervalle je Anlagentyp: nach welcher Regel wird heute terminiert", "biotec",
     "vor Konfiguration"],
    ["M6", "Bestätigung des gültigen Kundenstands – im Altsystem liegen mehrere Fassungen "
           "nebeneinander", "biotec", "vor Migration"],
    ["M7", "Freigabe der Belegmuster durch die Geschäftsführung, je Belegart", "biotec",
     "Meilenstein 2"],
    ["M8", "Benennung der vier Personen für die Vor-Ort-Schulung", "biotec",
     "vor Meilenstein 5"],
    ["M9", "Verfügbarkeit der Fachbereiche für Abstimmung, Test und Betriebssimulation – "
           "etwa ein halber Tag je Woche und Bereich, in der Simulationsphase mehr", "biotec",
     "laufend"],
    ["M10", "Zugang zum produktiven Bestand des Altsystems und eine Sicherung zum Testen, "
            "dazu Zugriff auf das Vorlagenverzeichnis", "Westbomke EDV", "Projektstart"],
    ["M11", "**Schriftliche Schnittstellenvereinbarung mit Westbomke EDV** mit benanntem "
            "Ansprechpartner und Reaktionszeit sowie deren Auftrag für die Arbeiten am "
            "Gutachtenprogramm", "biotec", "vor Meilenstein 2"],
    ["M12", "Inhaltliche Datenbereinigung nach unserer Vorlage: Kundenliste, "
            "Herstellerbezeichnungen, Wartungsfirmen. Reduziert Position 5.9 – und nur so "
            "stehen am Ende die fachlich richtigen Bezeichnungen im System", "biotec",
     "vor Migration"],
    ["M13", "Doppelerfassung während der Betriebssimulation: die Fälle laufen einen Monat in "
            "beiden Systemen", "biotec", "Simulationsphase"],
    ["M14", "Entscheidung über das Betriebsmodell von Odoo – Voraussetzung der Machbarkeit, "
            "siehe Annahmen", "biotec", "vor Beauftragung"],
]

MEILENSTEINE = [
    ["1", "Grobkonzept abgestimmt und freigegeben", "Konzeptdokument, Abnahmeprotokoll",
     "S − 26"],
    ["2", "Detailspezifikation freigegeben, Belegmuster freigegeben",
     "Spezifikation, Layoutmuster", "S − 20"],
    ["3", "Grundsystem konfiguriert, Stammdatenbrücke im Probebetrieb",
     "Testsystem, Brückenprotokoll", "S − 14"],
    ["4", "Daten übernommen, Summen abgestimmt", "Migrationsprotokoll", "S − 10"],
    ["5", "Abnahmetest bestanden, Betriebssimulation abgeschlossen",
     "Testprotokoll, Fehlerliste geschlossen", "S − 5"],
    ["6", "Schulung durchgeführt", "Teilnehmerliste, Anwenderleitfäden", "S − 2"],
    ["7", "**Betriebsbereitschaft** – Stichtag für die Aktivierung",
     "Abnahmeprotokoll Betriebsbereitschaft", "S"],
    ["8", "Betreuungsphase abgeschlossen, erster Monatsabschluss begleitet",
     "Abschlussbericht, offene Punkte übergeben", "S + 6"],
]

RISIKEN = [
    ["Arbeiten am Gutachtenprogramm", "Die Abschaltung der Erfassungsmasken und ein "
     "etwaiger Rückmeldekanal betreffen die Delphi-Anwendung. Diese Arbeiten erbringt "
     "Westbomke EDV auf eigenen Auftrag von biotec und sind hier nicht enthalten",
     "Schnittstellenvereinbarung M11, Rückfallebene als Option O7. Die Menüpunkte lassen "
     "sich voraussichtlich über die bestehende Rechtetabelle des Altsystems abschalten, "
     "ohne Eingriff in den Programmcode"],
    ["Altsystem ohne Dokumentation", "Alle Erkenntnisse zu Tabellen und Feldern stammen aus "
     "unserer eigenen Auswertung von Quellcode und Datenbestand, nicht aus einer "
     "Herstellerdokumentation", "Abweichende Feldbedeutungen werden nach Aufwand nachgezogen "
     "und vorher angezeigt"],
    ["Dauerhafter Betrieb zweier Systeme", "Es gibt kein Abschaltdatum für das "
     "Gutachtenprogramm. Die eingesetzte Datenbankversion MariaDB 10.4 erhält seit Juni 2024 "
     "keine Sicherheitsupdates mehr", "Feststellung, kein Gegenstand dieses Angebots. Ein "
     "Versionswechsel sollte unabhängig davon eingeplant werden"],
    ["Jahresrelease von Odoo", "Eigenmodule und Brücke müssen bei jedem Versionswechsel "
     "nachgezogen werden", "Option O13 Anwendungsbetreuung mit Versionspflege"],
    ["Verfügbarkeit der Fachbereiche", "Vier Personen tragen gleichzeitig den Betrieb. Bei "
     "Volllast mit 100 Anlagen je Woche ist der halbe Tag je Woche schwer zu halten",
     "Umstellungsstichtag außerhalb der Hochsaison, Meilensteine relativ zum Stichtag, "
     "Simulationsphase in eine ruhigere Periode legen"],
    ["Umstellungsstichtag", "Ein Stichtag mitten im Monat erschwert Saldenvortrag, Inventur "
     "und Umsatzsteuervoranmeldung", "Stichtag auf einen Monatsersten legen, Empfehlung "
     "1. Januar oder 1. April"],
    ["Datenqualität", "Die vorliegenden Bestände enthalten Testdatensätze, Dubletten und "
     "Freitextfelder. Der Bereinigungsaufwand ist nicht exakt vorhersehbar",
     "Position 5.9 und Mitwirkung M12; Abweichungen werden vor der Ausführung angezeigt"],
    ["Interessentenbestand und DSGVO", "10.387 Adressen ohne dokumentierte Rechtsgrundlage",
     "Prüfung vor Option O1, Herkunfts- und Rechtsgrundlagenfeld je Datensatz"],
]

def build(ziel):
    doc = Document()
    set_base_style(doc)
    seite_einrichten(doc, "Angebot Odoo-Einführung  ·  Dok.-Nr. " + DOK + "  ·  " + STAND)
    titel(doc, "Angebot: Einführung von Odoo", "biotec GmbH  ·  " + STAND)

    table(doc, ["", "Angaben"], [4.6, 12.4], [
        ["Angebot für", "biotec GmbH, Umwelt-Analytik-Beratung-Service\n"
                        "Elbrachtsweg 76, 33332 Gütersloh\n"
                        "z. Hd. Herrn Dr. Thomas Wilke, Geschäftsführer"],
        ["Ansprechpartner biotec", "Michael Brand, Projektleitung Hygieneinspektion "
                                   "(fachlich und technisch)\n"
                                   "Nicole Krupa (kaufmännisch und organisatorisch)"],
        ["Angebot von", "CertoClav Sterilizer GmbH, Geschäftsbereich Software\n"
                        "Peintner Straße 10, 4060 Leonding, Österreich\n"
                        "Michael Simon (geb. Dirix), MSc., Geschäftsführer\n"
                        "michael.simon@certoclav.com"],
        ["Dokument", DOK + "  ·  " + STAND],
        ["Gültigkeit", "60 Tage ab Ausstellungsdatum. Die kalkulierten Sätze gelten für "
                       "Leistungen innerhalb von zwölf Monaten nach Beauftragung"],
    ])
    doc.add_paragraph()

    heading(doc, "Ausgangslage", kicker="Worum es geht")
    absatz(doc, "biotec arbeitet an drei Standorten – Gütersloh mit Labor und Verwaltung, "
                "München mit den Hygienekontrollen und dem Technologiepark Mittweida für die "
                "Analytik von Boden und Wasser. Die kaufmännischen Abläufe laufen heute "
                "außerhalb eines Systems: Rechnungen, Buchhaltung, Artikel und Lager, "
                "Zeiterfassung und das gesamte Schulungsgeschäft werden von Hand geführt. "
                "Ein Buchhaltungssystem besteht nicht.")
    absatz(doc, "Die fachliche Arbeit trägt eine über Jahre gewachsene Eigenentwicklung. Wir "
                "haben Quellcode und Datenbestand ausgewertet: 255 Programmeinheiten, 430 "
                "Tabellen, Gutachten entstehen über eine automatisierte Satzstrecke. Rund die "
                "Hälfte der Anwendung deckt Aufgaben ab, die Odoo mitbringt. Die andere "
                "Hälfte ist Fachlogik ohne Entsprechung: Prüflisten nach VDI 6022, "
                "Labormasken, Maßnahmenkatalog und die Gutachtenerzeugung.")
    note(doc, "Daraus folgt der Zuschnitt: Odoo übernimmt das Kaufmännische, die bewährte "
              "Gutachtenerstellung bleibt. Verbunden werden beide über eine Stammdatenbrücke. "
              "Was funktioniert, wird nicht ersetzt.")

    heading(doc, "Der Zuschnitt", kicker="Zwei Systeme, eine Datenbasis")
    table(doc, ["Bereich", "Künftig", "Anmerkung"], [5.4, 3.4, 8.2], [
        ["Akquise und Interessenten", "**Odoo**", "10.387 Adressen, Option O1"],
        ["Kunden, Werke, Gebäude", "**Odoo**", "führender Stand für beide Systeme"],
        ["Anlagenstamm, Prüfintervalle", "**Odoo**", "938 RLT-Anlagen, Turnussteuerung neu"],
        ["Angebote und Aufträge", "**Odoo**", "Auftragskennung wird an das Altsystem übergeben"],
        ["Rechnungen und Buchhaltung", "**Odoo**", "heute kein System, SKR04"],
        ["Artikel, Lager, Einkauf", "**Odoo**", "heute mehrere Einzeldokumente"],
        ["Zeiterfassung", "**Odoo**", "heute außerhalb"],
        ["Schulungsgeschäft", "**Odoo**", "heute außerhalb, Phase 2"],
        ["Zusatzsoftware Anlagenliste", "**entfällt**", "Druckausgabe kommt aus Odoo"],
        ["Prüflisten nach VDI 6022", "bestehendes System", "bleibt unverändert"],
        ["Gutachten erstellen und bearbeiten", "bestehendes System", "Layout und Ablauf bleiben"],
        ["Labormasken, Messergebnisse", "bestehendes System", "bleibt unverändert"],
        ["Fotodokumentation vor Ort", "bestehendes System", "Odoo bietet keine Offline-Erfassung"],
    ], zeilen_fuellung=lambda z: BLAU if "Odoo" in z[1] else None)
    absatz(doc, "Die Stammdatenbrücke spiegelt Kunden und Anlagen aus Odoo in den Altbestand, "
                "sodass die Gutachtenanwendung sie unverändert liest. Die Erfassungsmasken für "
                "Kunden, Akquise und Angebote im Altsystem werden abgeschaltet, damit "
                "Stammdaten nur an einer Stelle entstehen – diese Arbeiten erbringt Westbomke "
                "EDV, siehe Risiken.", abstand=12)

    doc.add_page_break()
    heading(doc, "Leistungsumfang", kicker="Acht Arbeitspakete")
    absatz(doc, "Die Gliederung unterscheidet aktivierungsbezogene von laufenden Leistungen. "
                "Zuordnungsregel: **Aktivierungsbezogen** sind Leistungen zur technischen "
                "Herstellung und Inbetriebnahme bis zum Meilenstein Betriebsbereitschaft. "
                "**Aufwand** sind Beratung, Datenaufbereitung, Befähigung, Betriebsübergang "
                "und alle Leistungen nach Betriebsbereitschaft. Die bilanzielle Beurteilung "
                "trifft biotec mit Steuerberater und Wirtschaftsprüfer.", abstand=10)

    for name, positionen in WORKSTREAMS:
        s = sum(p[2] for p in positionen)
        heading(doc, name, kicker=pt(s) + " PT")
        zeilen = [[nr, inhalt, "Akt." if art == AKT else "Aufw.",
                   str(ph), pt(tage)] for nr, inhalt, tage, art, ph in positionen]
        zeilen.append(["", "**Summe " + name.split(" · ")[0] + "**", "", "",
                       "**" + pt(s) + "**"])
        table(doc, ["Nr.", "Leistung", "Einordn.", "Ph.", "PT"],
              [1.4, 11.6, 1.8, 0.9, 1.3], zeilen,
              zeilen_fuellung=lambda z: BLAU if z[0] == "" else None)
        doc.add_paragraph()

    # ---- Zahlen
    doc.add_page_break()
    heading(doc, "Preisübersicht", kicker="Vollumfang")
    zeilen = []
    gesamt = akt = 0.0
    for name, positionen in WORKSTREAMS:
        s = sum(p[2] for p in positionen)
        a = sum(p[2] for p in positionen if p[3] == AKT)
        gesamt += s
        akt += a
        zeilen.append([name, pt(a), pt(s - a), pt(s), eur(s * PT_SATZ)])
    zeilen.append(["**Gesamt**", "**" + pt(akt) + "**", "**" + pt(gesamt - akt) + "**",
                   "**" + pt(gesamt) + "**", "**" + eur(gesamt * PT_SATZ) + "**"])
    table(doc, ["Arbeitspaket", "Akt. PT", "Aufw. PT", "PT", "Betrag"],
          [7.4, 2.2, 2.4, 1.8, 3.2], zeilen,
          zeilen_fuellung=lambda z: GELB if z[1].startswith("**") else None)
    absatz(doc, f"Stundensatz {STUNDENSATZ:.2f} €".replace(".", ",") +
                f", Personentag zu {STD_JE_TAG} Stunden entspricht {eur(PT_SATZ)}. Alle "
                "Beträge netto zuzüglich Umsatzsteuer. Der Satz liegt im Marktband deutscher "
                "Odoo-Partner.", groesse=9, kursiv=True, abstand=10)
    absatz(doc, "**Aktivierungsbezogen " + pt(akt) + " PT (" + eur(akt * PT_SATZ) +
                "), Aufwand " + pt(gesamt - akt) + " PT (" + eur((gesamt - akt) * PT_SATZ) +
                ").**", abstand=12)

    # ---- Phasen
    heading(doc, "Empfohlener Zuschnitt in Phasen", kicker="Unser Vorschlag")
    absatz(doc, "Der Vollumfang lässt sich in einem Zug beauftragen. Wir empfehlen die "
                "Aufteilung in Phasen: Jede Phase erreicht eigene Betriebsbereitschaft, was "
                "die bilanzielle Beurteilung erleichtert, und die Belastung der Fachbereiche "
                "verteilt sich.")
    ph = {}
    for name, positionen in WORKSTREAMS:
        for nr, inhalt, tage, art, p_ in positionen:
            ph[p_] = ph.get(p_, 0.0) + tage
    table(doc, ["Phase", "Inhalt", "PT", "Betrag"], [2.6, 8.6, 2.0, 3.8], [
        ["**Phase 1**", "Kaufmännischer Kern, Anlagenstamm mit Turnussteuerung, "
                        "Stammdatenbrücke, Migration, Test, Schulung, Umstellung",
         pt(ph.get(1, 0)), eur(ph.get(1, 0) * PT_SATZ)],
        ["**Phase 2**", "Schulungsgeschäft vollständig – nach der Hochsaison, mit eigener "
                        "Betriebsbereitschaft", pt(ph.get(2, 0)), eur(ph.get(2, 0) * PT_SATZ)],
        ["**Phase 3**", "Gruppenthemen: LucaNet, IFRS-Überleitung, POC/WIP – sobald die "
                        "Vorgaben der Certania-Gruppe vorliegen (Optionen O2, O3, O5)",
         "nach Bedarf", "siehe Optionen"],
    ])
    absatz(doc, "In Phase 1 ist das Schulungsgeschäft nicht enthalten; es läuft bis dahin "
                "weiter wie heute.", groesse=9, kursiv=True, abstand=12)

    heading(doc, "Optionen", kicker="Auf Abruf, nicht im Vollumfang enthalten")
    zeilen = [[nr, inhalt, ("nach Abruf" if tage == 0 else pt(tage)),
               ("nach Abruf" if tage == 0 else eur(tage * PT_SATZ))]
              for nr, inhalt, tage in OPTIONEN]
    table(doc, ["Nr.", "Leistung", "PT", "Betrag"], [1.4, 10.6, 2.0, 3.0], zeilen)
    absatz(doc, "Optionen werden erst nach schriftlicher Freigabe erbracht. O2a und O2b sind "
                "Alternativen; welche gilt, entscheidet die Vorlage eines Lieferformats durch "
                "die Gruppe. O7 greift nur, wenn Westbomke EDV die Arbeiten am "
                "Gutachtenprogramm nicht übernimmt.", groesse=9, kursiv=True, abstand=12)

    # ---- Was entsteht
    doc.add_page_break()
    heading(doc, "Was bei biotec entsteht", kicker="Arbeitsergebnisse")
    absatz(doc, "Aus dem Projekt gehen benannte Arbeitsergebnisse in das Eigentum von biotec "
                "über. Sie werden je Arbeitspaket in einem Übergabeverzeichnis dokumentiert "
                "(Position 1.8):")
    bullets(doc, [
        "**Erweiterungsmodul `biotec_rlt`** – RLT-Anlagenstamm mit VDI-Feldern, Kennungen, "
        "Standortstruktur und Prüfhistorie, als Quellcode",
        "**Stammdatenbrücke** zum Gutachtenprogramm, als Quellcode samt Betriebsanleitung",
        "**Belegvorlagen** für sieben Belegarten im biotec-Layout",
        "**Migrationsprogramme** mit Feldzuordnung und Wiederholbarkeit",
        "**Auswertungen und Kennzahlenberichte**",
        "**Dokumentation:** Detailspezifikation, Testprotokolle, "
        "GoBD-Verfahrensdokumentation, Migrationsdokumentation, Betriebshandbuch",
    ])
    absatz(doc, "Der Quellcode wird in einem Repository von biotec abgelegt. biotec erhält "
                "ein unbefristetes, übertragbares Nutzungs- und Bearbeitungsrecht daran.",
           abstand=12)

    heading(doc, "Layouttreue – was darunter zu verstehen ist", kicker="Abnahmekriterium")
    absatz(doc, "Layouttreue heißt: gleiche Gestaltungselemente, gleiche Reihenfolge, gleiche "
                "Pflichtangaben, Briefpapier und Schriftbild wie bisher. Die Abnahme erfolgt "
                "anhand eines je Belegart freigegebenen Musters (Meilenstein 2). Eine "
                "pixelgenaue Übereinstimmung mit der bisherigen Satzstrecke ist technisch "
                "nicht herstellbar und nicht Gegenstand der Abnahme. Nach Musterfreigabe sind "
                "zwei Korrekturschleifen je Belegart enthalten.", abstand=12)

    heading(doc, "Zeitplan und Meilensteine", kicker="Rund 32 Wochen")
    absatz(doc, "Die Angaben sind Wochen relativ zum Umstellungsstichtag **S**. Der Stichtag "
                "sollte auf einem Monatsersten außerhalb der Hochsaison liegen – empfohlen "
                "1. Januar oder 1. April.")
    table(doc, ["Nr.", "Meilenstein", "Nachweis", "Woche"], [1.2, 7.0, 6.0, 2.8],
          MEILENSTEINE, zeilen_fuellung=lambda z: GELB if z[0] == "7" else None)
    absatz(doc, "Der Meilenstein Betriebsbereitschaft ist der Stichtag, bis zu dem "
                "zurechenbare Herstellungskosten anfallen. Er wird mit einem eigenen "
                "Abnahmeprotokoll dokumentiert.", groesse=9, kursiv=True, abstand=12)

    # ---- Mitwirkung
    heading(doc, "Mitwirkung von biotec", kicker="Terminkritisch")
    absatz(doc, "Eine Einführung gelingt nur gemeinsam. Die folgenden Punkte liegen bei "
                "biotec. Ohne sie verschiebt sich der Plan.")
    table(doc, ["Nr.", "Aufgabe", "Wer", "Wann"], [1.2, 11.0, 3.0, 1.8], MITWIRKUNG,
          zeilen_fuellung=lambda z: GELB if z[0] in ("M1", "M2", "M11", "M14") else None)

    # ---- Risiken
    doc.add_page_break()
    heading(doc, "Risiken und wie wir damit umgehen", kicker="Offen benannt")
    table(doc, ["Thema", "Sachverhalt", "Umgang"], [3.6, 6.7, 6.7], RISIKEN)

    heading(doc, "Annahmen und Voraussetzungen", kicker="Woran die Kalkulation hängt")
    bullets(doc, [
        "**Betriebsmodell als Vorfrage:** Odoo.sh oder Betrieb auf eigener beziehungsweise "
        "dedizierter Infrastruktur ist technische Voraussetzung der Positionen 3.1 bis 3.5 "
        "und 4.1. Auf Odoo Online sind eigene Module nicht ausführbar; in diesem Fall ändern "
        "sich Umfang, Aufwand und die bilanzielle Beurteilung. Diese Entscheidung sollte vor "
        "der Beauftragung fallen.",
        "Die drei Standorte werden als eine Gesellschaft mit analytischer Trennung "
        "abgebildet, nicht als getrennte Mandanten.",
        "Die Buchführung bleibt bei biotec; der Steuerberater erhält einen DATEV-Export. Die "
        "Übermittlung der Umsatzsteuervoranmeldung an ELSTER erfolgt außerhalb von Odoo.",
        "Das Gutachtenlayout bleibt unverändert und wird nicht in Odoo nachgebaut.",
        "Die Prüflisten nach VDI 6022 bleiben im bestehenden System.",
        "Der Kundenstamm wird neu aufgebaut, nicht eins zu eins übernommen. Die vorliegende "
        "Liste ist eine Kontaktliste ohne Kundennummern und Konditionen.",
        "Lagerbestände werden zum Stichtag gezählt und nicht aus dem Altsystem übernommen.",
        "Die umsatzsteuerliche Behandlung der VDI-Schulungen ist zu klären. Kalkuliert ist "
        "die Regelbesteuerung; eine Steuerbefreiung nach § 4 UStG würde Erlöskonten und "
        "Belegtexte verändern.",
        "Zeitaufnahmen der heutigen Abläufe liegen nicht vor. Die Aufwände beruhen auf der "
        "Auswertung der gelieferten Daten und auf Erfahrungswerten.",
        "**Bilanzielle Beurteilung:** Die Behandlung im HGB-Einzelabschluss und im "
        "IFRS-Konzernabschluss der Gruppe können voneinander abweichen. Bei Betrieb als reine "
        "Mietlösung ist im IFRS-Abschluss überwiegend Aufwand zu erwarten. Dieses Angebot "
        "liefert die Grundlage für beide Beurteilungen; die Beurteilung selbst trifft biotec "
        "mit Steuerberater und Wirtschaftsprüfer.",
        "**Gesellschafterverhältnisse:** Die CertoClav Sterilizer GmbH steht im Eigentum von "
        "Michael Simon und Moritz Gruber. Moritz Gruber ist zugleich in der Certania-Gruppe "
        "tätig. Wir legen das offen, damit die Beauftragung als Geschäft mit nahestehenden "
        "Personen behandelt und der Fremdvergleich geführt werden kann.",
    ])

    heading(doc, "Nicht enthalten", kicker="Klarstellung")
    bullets(doc, [
        "Lizenz-, Hosting- und Betriebskosten für Odoo sowie Hardware",
        "Änderungen am Programmcode der Gutachtenanwendung. Die Abschaltung der "
        "Erfassungsmasken und ein etwaiger Rückmeldekanal auf Seiten des Altsystems erbringt "
        "Westbomke EDV auf eigenen Auftrag von biotec – ersatzweise Option O7",
        "Abbildung der Prüflisten, der Labormasken oder der Gutachtenerzeugung in Odoo",
        "**Fotodokumentation und Prüflistenerfassung vor Ort.** Odoo bietet keine "
        "Offline-Erfassung; eine mobile Erfassung in Technikzentralen ohne Netzabdeckung ist "
        "nicht Gegenstand dieses Angebots",
        "**Umsatzrealisierung nach Fertigstellungsgrad (POC), WIP-Bewertung, "
        "Cost-to-Complete und automatische Abgrenzung von Fremdleistungen.** Odoo bietet "
        "dafür keine Standardfunktion; die vorhandene Rechnungsabgrenzung ist zeit- und nicht "
        "fortschrittsbasiert. Adressierbar über die Optionen O5a und O5b",
        "Webshop und Online-Buchung der Seminare – Option O6",
        "Fertigung und Stücklisten über die vorhandene geringe Eigenproduktion hinaus",
        "Anwendungsbetreuung nach Abschluss der Betreuungsphase – Option O13",
        "Steuerliche, bilanzielle und rechtliche Beratung",
    ])

    heading(doc, "Vertragliche Rahmenbedingungen", kicker="Kurzfassung")
    bullets(doc, [
        "Abgerechnet wird nach Leistungsfortschritt je Arbeitspaket. Jeder Rechnung liegt ein "
        "Leistungsnachweis mit Leistungsbeschreibung **und Zeitaufstellung je Arbeitspaket** "
        "bei – das ist die Grundlage für die Zuordnung der Herstellungskosten.",
        "Rechnungsstellung monatlich, Zahlungsziel 21 Tage netto.",
        "Abweichungen von mehr als zehn Prozent je Arbeitspaket werden vor der Ausführung "
        "angezeigt und schriftlich freigegeben. Die genannten Beträge sind die kalkulierte "
        "Grundlage, kein Festpreis.",
        "Gewährleistung: sechs Monate ab Betriebsbereitschaft auf die vereinbarte "
        "Funktionsfähigkeit der Arbeitsergebnisse. Mängel werden nachgebessert.",
        "Haftung: für Vorsatz und grobe Fahrlässigkeit unbeschränkt; im Übrigen begrenzt auf "
        "die Höhe des Auftragswerts. Keine Haftung für Datenverlust, soweit biotec die "
        "vereinbarte Datensicherung nicht vorhält.",
        "Nutzungsrechte an den Arbeitsergebnissen wie im Abschnitt Arbeitsergebnisse "
        "beschrieben. Vorbestehende Werkzeuge und Bibliotheken von CertoClav bleiben deren "
        "Eigentum; biotec erhält daran ein einfaches Nutzungsrecht für den Betrieb.",
        "Verzögert sich eine Mitwirkungsleistung nach dem Abschnitt Mitwirkung, verschieben "
        "sich die davon abhängigen Termine entsprechend; Wartezeiten werden nach Aufwand "
        "abgerechnet.",
        "Reisekosten nach Aufwand, vorab abgestimmt.",
        "Beide Seiten können den Auftrag mit vier Wochen Frist zum Ende eines Arbeitspakets "
        "kündigen. Erbrachte Leistungen werden abgerechnet, Arbeitsergebnisse und "
        "Dokumentation übergeben.",
        "Es gilt deutsches Recht, Gerichtsstand Gütersloh. Vertraulichkeit und "
        "Datenverarbeitung richten sich nach NDA-2026-001 und AVV-2026-001.",
    ])

    heading(doc, "Anhang: Lizenzbedarf, informativ", kicker="Nicht Teil des Angebots")
    absatz(doc, "Zur Einordnung der Gesamtkosten, ohne Anspruch auf Genauigkeit: Im Altsystem "
                "bestehen 17 Benutzerkonten, davon zwei geteilte Laborkonten und zwei "
                "Praktikantenkonten. Für Odoo ist von etwa 12 bis 15 benannten Nutzern "
                "auszugehen. Wegen der eigenen Module ist die Enterprise-Edition mit einem "
                "Betriebsmodell erforderlich, das eigene Module zulässt. Lizenzpreise bezieht "
                "biotec direkt bei Odoo; wir beraten bei der Auswahl, verdienen daran nicht.",
           abstand=12)

    heading(doc, "Nächster Schritt", kicker="")
    absatz(doc, "Wir schlagen vor, dieses Angebot in einem Termin gemeinsam durchzugehen, den "
                "Zuschnitt in Phasen zu entscheiden und die Punkte unter Mitwirkung zu "
                "terminieren. Vorab zu klären sind das Betriebsmodell von Odoo (M14) und der "
                "Auftrag an Westbomke EDV (M11).")
    absatz(doc, "Ergänzend stellen wir eine Betrachtung der freiwerdenden Personalkapazität "
                "und des Einsparpotenzials zur Verfügung – als eigenes Dokument, damit "
                "Kosten und Nutzen getrennt beurteilt werden können.", abstand=18)

    table(doc, ["biotec GmbH", "CertoClav Sterilizer GmbH"], [8.5, 8.5], [
        ["\n\nOrt, Datum\n\n\n\n\n________________________________\n"
         "Dr. Thomas Wilke\nGeschäftsführer",
         "\n\nOrt, Datum\n\n\n\n\n________________________________\n"
         "Michael Simon (geb. Dirix), MSc.\nGeschäftsführer"],
    ])

    doc.save(ziel)
    so = sum(o[2] for o in OPTIONEN)
    print("geschrieben:", ziel)
    print(f"Vollumfang: {pt(gesamt)} PT = {eur(gesamt*PT_SATZ)}")
    print(f"  aktivierungsbezogen: {pt(akt)} PT = {eur(akt*PT_SATZ)} ({akt/gesamt*100:.0f} %)")
    print(f"  Aufwand:             {pt(gesamt-akt)} PT = {eur((gesamt-akt)*PT_SATZ)}")
    for p_ in sorted(ph):
        print(f"  Phase {p_}: {pt(ph[p_])} PT = {eur(ph[p_]*PT_SATZ)}")
    print(f"Optionen (ohne O13): {pt(so)} PT = {eur(so*PT_SATZ)}")


if __name__ == "__main__":
    build(sys.argv[1] if len(sys.argv) > 1 else "Angebot_Odoo_biotec.docx")
