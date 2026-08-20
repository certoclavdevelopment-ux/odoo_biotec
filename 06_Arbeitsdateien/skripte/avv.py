#!/usr/bin/env python3
"""Erzeugt den Vertrag über die Auftragsverarbeitung (AVV) nach Art. 28 DSGVO.

Ergänzt die Vertraulichkeitsvereinbarung (NDA-2026-001, Ziffer 5.2). Verantwortlicher
ist biotec, Auftragsverarbeiter CertoClav.

Unterschriftsfertig, keine Platzhalter. Ueber den Schalter KI_ERLAUBT wird
festgelegt, ob KI-Dienste personenbezogene Daten verarbeiten duerfen:

  KI_ERLAUBT = True  (Standard) – KI-Anbieter sind in Anlage 2 benannt, Ziffer 4.4
    erlaubt den Einsatz, Ziffer 4.5 fordert kein Modelltraining, geschaeftliche
    Nutzungsbedingungen mit AV-Vertrag und Datenminimierung.
  KI_ERLAUBT = False – Zusage, dass keine personenbezogenen Daten an KI-Dienste
    gehen. Enger, aber ohne Drittlandthematik.

Geprueft am 20.08.2026: Anthropic stellt einen AV-Vertrag (DPA) bereit, der
automatisch Bestandteil der Commercial Terms ist und keine separate Unterschrift
braucht. Er gilt fuer Claude for Work (Team/Enterprise) und die API, nicht fuer
Free/Pro. Enthalten sind die EU-Standardvertragsklauseln Modul 2 und 3. Eine
EU-Datenresidenz ist fuer den Team-Tarif nicht verfuegbar - Verarbeitung in den USA.

Aufruf:  python3 avv.py <zieldatei.docx>
"""
import sys

from docx import Document
from docx.shared import Cm, Pt

from docx_bausteine import ACCENT, GREY, heading, seite_einrichten, table, titel

STAND = "Version 2, Stand 20.08.2026"

# KI_ERLAUBT = True  -> KI-Dienste sind als Unterauftragsverarbeiter benannt und
#                       duerfen personenbezogene Daten verarbeiten (Ziffer 4.4 offen).
# KI_ERLAUBT = False -> Zusage, dass keine personenbezogenen Daten an KI-Dienste gehen.
#                       Enger, aber ohne Drittlandthematik.
KI_ERLAUBT = True


def para(doc, text, groesse=10, abstand=6, kursiv=False, einzug=0.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(abstand)
    if einzug:
        p.paragraph_format.left_indent = Cm(einzug)
    for i, teil in enumerate(text.split("**")):
        if not teil:
            continue
        r = p.add_run(teil)
        r.font.size = Pt(groesse)
        r.font.italic = kursiv
        r.font.bold = i % 2 == 1
    return p


def ziffer(doc, nummer, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(1.1)
    p.paragraph_format.first_line_indent = Cm(-1.1)
    r = p.add_run(nummer + "\t")
    r.font.size = Pt(10)
    r.font.bold = True
    for i, teil in enumerate(text.split("**")):
        if not teil:
            continue
        rr = p.add_run(teil)
        rr.font.size = Pt(10)
        rr.font.bold = i % 2 == 1


def punkte(doc, items):
    for t in items:
        para(doc, "– " + t, groesse=10, abstand=2, einzug=1.6)


def paragraf(doc, nummer, titel_text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    r = h.add_run(f"{nummer}. {titel_text}")
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = ACCENT


def build(path):
    doc = Document()
    seite_einrichten(doc, "Auftragsverarbeitung  ·  Dok.-Nr. AVV-2026-001  ·  " + STAND)
    titel(doc, "Vertrag über die Auftragsverarbeitung",
          "nach Art. 28 DSGVO  ·  " + STAND)

    para(doc, "zwischen", abstand=4)
    table(doc, ["Partei", "Angaben"], [4.2, 12.8], [
        ["biotec GmbH\n**Verantwortlicher**",
         "Elbrachtsweg 76, 33332 Gütersloh, Deutschland\n"
         "Telefon +49 5241 307 20-0 · info@biotec-gmbh.com\n"
         "Amtsgericht Gütersloh, HRB 3829\n"
         "vertreten durch: Dr. Andreas Bermpohl, Prokurist"],
        ["CertoClav Sterilizer GmbH\n**Auftragsverarbeiter**",
         "Peintner Straße 10, 4060 Leonding, Österreich\n"
         "Telefon +43 732 674 278 · support@certoclav.com\n"
         "Firmenbuch: Landesgericht Linz, FN 122912d · UID ATU22821702\n"
         "vertreten durch: Michael Simon (geb. Dirix), MSc., Geschäftsführer"],
    ])
    para(doc, "Dieser Vertrag ergänzt die Vertraulichkeitsvereinbarung Dok.-Nr. NDA-2026-001 "
              "und konkretisiert deren Ziffer 5.2.", groesse=9.5, kursiv=True, abstand=10)

    paragraf(doc, 1, "Gegenstand, Dauer und Weisungsbindung")
    ziffer(doc, "1.1", "Gegenstand ist die Verarbeitung personenbezogener Daten durch den "
                       "Auftragsverarbeiter im Rahmen der Einführung des ERP-Systems Odoo bei "
                       "biotec – insbesondere Analyse der Altdaten, Datenmigration, "
                       "Konfiguration, Test und Schulungsvorbereitung.")
    ziffer(doc, "1.2", "Die Verarbeitung erfolgt ausschließlich **auf dokumentierte Weisung** "
                       "des Verantwortlichen und ausschließlich zu diesem Zweck. Eine "
                       "Verarbeitung für eigene Zwecke findet nicht statt.")
    ziffer(doc, "1.3", "Der Vertrag beginnt mit Unterzeichnung und endet mit Abschluss des "
                       "Projekts, spätestens mit Löschung der Daten nach Ziffer 9. Er kann von "
                       "beiden Seiten jederzeit gekündigt werden.")
    if KI_ERLAUBT:
        ziffer(doc, "1.4", "Die Verarbeitung findet grundsätzlich in der Europäischen Union bzw. "
                           "dem Europäischen Wirtschaftsraum statt. Übermittlungen in Drittländer "
                           "erfolgen ausschließlich an die in **Anlage 2** benannten Empfänger "
                           "und nur auf Grundlage eines Angemessenheitsbeschlusses der "
                           "Europäischen Kommission oder geeigneter Garantien nach "
                           "Art. 46 DSGVO, insbesondere der Standardvertragsklauseln.")
    else:
        ziffer(doc, "1.4", "Die Verarbeitung findet **ausschließlich in der Europäischen Union "
                           "bzw. dem Europäischen Wirtschaftsraum** statt. Eine Übermittlung in "
                           "Drittländer erfolgt nicht.")

    ziffer(doc, "1.5", "**Weisungsberechtigt** auf Seiten des Verantwortlichen sind seine "
                       "gesetzlichen Vertreter und Prokuristen sowie die von biotec benannten "
                       "Projektansprechpartner, derzeit Nicole Krupa und Michael Brand. "
                       "Weisungen nimmt auf Seiten des Auftragsverarbeiters Michael Simon "
                       "entgegen. Weisungen werden in Textform erteilt; mündlich erteilte "
                       "Weisungen bestätigt der Auftragsverarbeiter in Textform.")

    paragraf(doc, 2, "Art der Daten und Kategorien betroffener Personen")
    ziffer(doc, "2.1", "Verarbeitet werden voraussichtlich folgende Datenarten:")
    punkte(doc, [
        "Kontaktdaten von Ansprechpartnern bei Kunden und Lieferanten: Name, Funktion, "
        "Anschrift, Telefon, E-Mail",
        "Daten zu Aufträgen, Anlagen, Prüfungen und Gutachten, soweit einer Person zuordenbar",
        "Teilnehmerdaten aus Schulungen: Name, Unternehmen, Kursteilnahme, Bescheinigungen",
        "Beschäftigtendaten in dem Umfang, der für Odoo-Nutzerkonten und Zeiterfassung "
        "erforderlich ist: Name, Rolle, Standort, dienstliche Kontaktdaten",
        "Zahlungs- und Abrechnungsdaten, soweit personenbezogen",
    ])
    ziffer(doc, "2.2", "Kategorien betroffener Personen: Ansprechpartner bei Kunden und "
                       "Lieferanten, Schulungsteilnehmende, Beschäftigte von biotec.")
    ziffer(doc, "2.3", "**Besondere Kategorien personenbezogener Daten nach Art. 9 DSGVO sind "
                       "nicht Gegenstand dieses Vertrags.** Der Verantwortliche stellt sicher, "
                       "dass solche Daten – etwa Gesundheits- oder Personalaktendaten – nicht "
                       "übermittelt werden, solange die Parteien hierzu nichts Abweichendes "
                       "schriftlich vereinbaren.")
    ziffer(doc, "2.4", "**Nicht personenbezogene Daten sind nicht Gegenstand dieses Vertrags.** "
                       "Dazu gehören insbesondere technische Anlagen- und Objektdaten, Artikel- "
                       "und Materialstamm, Kontenrahmen und Buchungsdaten ohne Personenbezug, "
                       "Prüfpläne und Analysemethoden sowie Datenbankstrukturen und Quellcode "
                       "der Altanwendung. Für diese gilt die Vertraulichkeitsvereinbarung "
                       "NDA-2026-001.")

    paragraf(doc, 3, "Pflichten des Auftragsverarbeiters")
    ziffer(doc, "3.1", "Der Auftragsverarbeiter verarbeitet die Daten nur im Rahmen des Auftrags "
                       "und dieses Vertrags. Hält er eine Weisung für rechtswidrig, teilt er "
                       "dies unverzüglich mit und darf deren Ausführung aussetzen.")
    ziffer(doc, "3.2", "Zum Zugriff berechtigt sind nur Personen, die die Daten für den Zweck "
                       "benötigen. Diese Personen sind zur Vertraulichkeit verpflichtet; die "
                       "Verpflichtung besteht über das Ende ihrer Tätigkeit hinaus.")
    ziffer(doc, "3.3", "Der Auftragsverarbeiter unterhält die in Anlage 1 beschriebenen "
                       "technischen und organisatorischen Maßnahmen nach Art. 32 DSGVO und passt "
                       "sie bei Bedarf an, ohne das Schutzniveau zu senken.")
    ziffer(doc, "3.4", "Er unterstützt den Verantwortlichen bei der Beantwortung von Anfragen "
                       "betroffener Personen nach Art. 15 bis 22 DSGVO sowie bei "
                       "Datenschutz-Folgenabschätzungen und Meldungen an Aufsichtsbehörden.")
    ziffer(doc, "3.5", "**Verletzungen des Schutzes personenbezogener Daten** meldet er dem "
                       "Verantwortlichen unverzüglich, spätestens innerhalb von **24 Stunden** "
                       "nach Kenntnis, mit Beschreibung des Vorfalls, der betroffenen "
                       "Datenkategorien und der getroffenen Maßnahmen.")
    ziffer(doc, "3.6", "Er führt ein Verzeichnis der im Auftrag durchgeführten "
                       "Verarbeitungstätigkeiten und benennt eine Kontaktstelle für "
                       "Datenschutzfragen: michael.simon@certoclav.com.")

    doc.add_page_break()

    paragraf(doc, 4, "Unterauftragsverarbeiter")
    ziffer(doc, "4.1", "Der Verantwortliche genehmigt die in **Anlage 2** aufgeführten "
                       "Unterauftragsverarbeiter.")
    ziffer(doc, "4.2", "Weitere Unterauftragsverarbeiter werden dem Verantwortlichen vorab "
                       "schriftlich mitgeteilt. Er kann innerhalb von **14 Tagen** widersprechen; "
                       "im Fall eines Widerspruchs suchen die Parteien eine Lösung, andernfalls "
                       "kann der betroffene Leistungsteil gekündigt werden.")
    ziffer(doc, "4.3", "Der Auftragsverarbeiter verpflichtet jeden Unterauftragsverarbeiter "
                       "schriftlich auf ein Schutzniveau, das diesem Vertrag mindestens "
                       "entspricht, und bleibt gegenüber dem Verantwortlichen verantwortlich.")
    if KI_ERLAUBT:
        ziffer(doc, "4.4", "Der Auftragsverarbeiter setzt **KI-gestützte Dienste** zur Analyse, "
                           "Aufbereitung und Migration der Daten sowie zur Konfiguration des "
                           "Zielsystems ein. Die eingesetzten Anbieter sind in Anlage 2 "
                           "benannt. Für sie gelten die Anforderungen der Ziffern 4.1 bis 4.3 "
                           "unverändert.")
        ziffer(doc, "4.5", "Für KI-gestützte Dienste stellt der Auftragsverarbeiter zusätzlich "
                           "sicher, dass")
        for t in ("die Daten **nicht zum Training** oder zur Verbesserung von Modellen "
                  "verwendet werden,",
                  "geschäftliche Nutzungsbedingungen mit Auftragsverarbeitungsvertrag "
                  "zugrunde liegen, keine Endkunden- oder Privatnutzungstarife,",
                  "Eingaben und Ausgaben nicht länger gespeichert werden als für die "
                  "Erbringung der Leistung erforderlich,",
                  "nur die für den jeweiligen Arbeitsschritt erforderlichen Daten übermittelt "
                  "werden – bei Strukturarbeiten ohne Personenbezug, bei Datenprüfungen "
                  "beschränkt auf den geprüften Ausschnitt."):
            para(doc, "– " + t, groesse=10, abstand=2, einzug=1.6)
    else:
        ziffer(doc, "4.4", "**Eine Verarbeitung personenbezogener Daten durch KI-Dienste findet "
                           "nicht statt.** Soweit der Auftragsverarbeiter KI-gestützte Werkzeuge "
                           "zur Erstellung von Konfigurationen, Feldzuordnungen und "
                           "Auswertungslogik einsetzt, geschieht dies ausschließlich anhand von "
                           "Strukturinformationen – Feldbezeichnungen, Datentypen, "
                           "Formatbeschreibungen – ohne personenbezogene Daten.")

    paragraf(doc, 5, "Kontrollrechte")
    ziffer(doc, "5.1", "Der Verantwortliche darf die Einhaltung dieses Vertrags überprüfen – "
                       "durch Auskunftsverlangen, Vorlage von Nachweisen oder, nach "
                       "angemessener Vorankündigung, durch Prüfung vor Ort während der "
                       "Geschäftszeiten.")
    ziffer(doc, "5.2", "Der Auftragsverarbeiter wirkt bei Prüfungen mit und stellt die "
                       "erforderlichen Auskünfte und Nachweise bereit.")

    paragraf(doc, 6, "Ort der Verarbeitung und eingesetzte Systeme")
    ziffer(doc, "6.1", "Die Daten werden auf zugriffsbeschränkten Systemen des "
                       "Auftragsverarbeiters sowie in den in Anlage 2 genannten Diensten "
                       "verarbeitet.")
    ziffer(doc, "6.2", "Datenbanksicherungen und der Quellcode der Altanwendung werden getrennt "
                       "von der übrigen Projektdokumentation und ohne Ablage in einem "
                       "Versionsverwaltungssystem aufbewahrt.")

    paragraf(doc, 7, "Berichtigung, Einschränkung, Löschung")
    ziffer(doc, "7.1", "Der Auftragsverarbeiter berichtigt, löscht oder schränkt die "
                       "Verarbeitung nur auf Weisung des Verantwortlichen ein.")
    ziffer(doc, "7.2", "Auskunft an betroffene Personen erteilt er nicht selbst, sondern "
                       "verweist an den Verantwortlichen und unterstützt ihn.")

    paragraf(doc, 8, "Haftung")
    ziffer(doc, "8.1", "Es gilt Art. 82 DSGVO. Im Übrigen richtet sich die Haftung nach den "
                       "Vereinbarungen des zugrunde liegenden Vertragsverhältnisses, soweit ein "
                       "solches besteht, andernfalls nach den gesetzlichen Vorschriften.")

    paragraf(doc, 9, "Beendigung, Löschung und Rückgabe")
    ziffer(doc, "9.1", "Nach Abschluss des Projekts – oder auf früheres Verlangen – gibt der "
                       "Auftragsverarbeiter alle personenbezogenen Daten zurück oder löscht sie "
                       "einschließlich Kopien und bestätigt dies schriftlich.")
    ziffer(doc, "9.2", "Ausgenommen sind Daten, für die eine gesetzliche Aufbewahrungspflicht "
                       "besteht, sowie Sicherungskopien aus automatisierten Backup-Verfahren, "
                       "die nach den regulären Aufbewahrungsfristen auslaufen. Für diese gelten "
                       "die Pflichten dieses Vertrags fort.")

    paragraf(doc, 10, "Schlussbestimmungen")
    ziffer(doc, "10.1", "Änderungen bedürfen der Schriftform. Bei Widersprüchen zwischen diesem "
                        "Vertrag und der Vertraulichkeitsvereinbarung gilt für den Datenschutz "
                        "dieser Vertrag.")
    ziffer(doc, "10.2", "Sollte eine Bestimmung unwirksam sein, bleibt der Vertrag im Übrigen "
                        "wirksam.")
    ziffer(doc, "10.3", "Es gilt das Recht der Bundesrepublik Deutschland. Ausschließlicher "
                        "Gerichtsstand ist Gütersloh.")

    heading(doc, "Unterschriften")
    table(doc, ["biotec GmbH – Verantwortlicher", "CertoClav Sterilizer GmbH – Auftragsverarbeiter"],
          [8.5, 8.5], [
        ["\n\nOrt, Datum\n\n\n\n\n________________________________\n"
         "Dr. Andreas Bermpohl\nProkurist",
         "\n\nOrt, Datum\n\n\n\n\n________________________________\n"
         "Michael Simon (geb. Dirix), MSc.\nGeschäftsführer"],
    ])

    doc.add_page_break()

    heading(doc, "Anlage 1 · Technische und organisatorische Maßnahmen",
            kicker="Art. 32 DSGVO")
    table(doc, ["Bereich", "Maßnahmen"], [4.6, 12.4], [
        ["Zutrittskontrolle", "Geschäftsräume mit Zutrittsbeschränkung; Serverräume nur für "
                              "berechtigtes Personal"],
        ["Zugangskontrolle", "Persönliche Benutzerkonten, keine gemeinsam genutzten Zugänge; "
                             "Mehrfaktor-Authentifizierung für Cloud-Dienste; "
                             "Bildschirmsperre; verschlüsselte Festplatten der Endgeräte"],
        ["Zugriffskontrolle", "Rechtevergabe nach dem Prinzip der geringsten Rechte; Zugriff "
                              "nur für Projektbeteiligte; Service-Benutzer ohne Recht zur "
                              "Finalbuchung"],
        ["Weitergabekontrolle", "Übertragung ausschließlich verschlüsselt (TLS); Austausch über "
                                "einen zugriffsbeschränkten Projektordner statt per "
                                "E-Mail-Anhang; keine Ablage von Kundendaten in "
                                "Versionsverwaltungssystemen"],
        ["Eingabekontrolle", "Protokollierung von Änderungen in Odoo; Lieferprotokoll je "
                             "Datenlieferung; Nachvollziehbarkeit der Migrationsschritte über "
                             "dokumentierte Feldzuordnungen"],
        ["Verfügbarkeitskontrolle", "Regelmäßige Sicherungen; Wiederherstellung getestet; "
                                    "Virenschutz und Firewall; Aktualisierung der Systeme"],
        ["Trennungskontrolle", "Projektdaten getrennt von anderen Mandanten und von der "
                               "eigenen Produktivumgebung von CertoClav; Test- und "
                               "Produktivumgebung getrennt"],
        ["Löschkonzept", "Löschung nach Projektabschluss gemäß Ziffer 9; "
                         "Datenbanksicherungen und Quellcode gesondert und nachweislich"],
        ["Organisation", "Verpflichtung der Mitarbeitenden auf Vertraulichkeit; Kontaktstelle "
                         "für Datenschutzfragen benannt; Verzeichnis der "
                         "Verarbeitungstätigkeiten"],
    ])

    heading(doc, "Anlage 2 · Unterauftragsverarbeiter")
    para(doc, "Der Verantwortliche genehmigt mit Unterzeichnung die folgenden "
              "Unterauftragsverarbeiter.", groesse=10)

    zeilen = [
        ["Microsoft Ireland Operations Ltd.\nOne Microsoft Place, Dublin, Irland",
         "OneDrive und SharePoint – Austausch und Ablage der Projektunterlagen",
         "EU/EWR",
         "Auftragsverarbeitungsvertrag von Microsoft (Data Protection Addendum), "
         "EU-Datengrenze"],
        ["Odoo S.A.\nChaussée de Namur 40, 1367 Grand-Rosière, Belgien",
         "Betrieb und Hosting der Odoo-Umgebung einschließlich der darin enthaltenen "
         "KI-Funktionen",
         "EU/EWR (Belgien)",
         "Auftragsverarbeitungsvertrag von Odoo S.A. samt dessen Unterauftragsverarbeitern"],
    ]
    if KI_ERLAUBT:
        zeilen.append(
            ["Anthropic PBC\nSan Francisco, Kalifornien, USA\n(Claude for Work – Team-Tarif)",
             "KI-gestützte Analyse, Aufbereitung und Migration der Daten sowie Konfiguration "
             "des Zielsystems",
             "USA\n(eine EU-Datenresidenz ist für diesen Tarif nicht verfügbar)",
             "Anthropic Data Processing Addendum, automatisch Bestandteil der Commercial "
             "Terms of Service; EU-Standardvertragsklauseln Modul 2 und 3 nach "
             "Durchführungsbeschluss (EU) 2021/914; bei Claude for Work keine Nutzung der "
             "Eingaben zum Modelltraining"])
    table(doc, ["Dienstleister", "Leistung", "Ort der Verarbeitung", "Grundlage"],
          [4.4, 4.8, 3.0, 4.8], zeilen)

    if not KI_ERLAUBT:
        para(doc, "**Nicht als Unterauftragsverarbeiter eingesetzt:** KI-Dienste. "
                  "Personenbezogene Daten werden nicht durch KI-Dienste verarbeitet – "
                  "siehe Ziffer 4.4.", groesse=10)
    para(doc, "Wird die Odoo-Umgebung auf eigener Infrastruktur von biotec betrieben, entfällt "
              "die Zeile zu Odoo S.A. Weitere Unterauftragsverarbeiter kommen nur nach dem "
              "Verfahren in Ziffer 4.2 hinzu.", groesse=9.5, kursiv=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run("Kontaktstelle Datenschutz beim Auftragsverarbeiter: Michael Simon, "
                  "michael.simon@certoclav.com, +43 732 674 278")
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY

    doc.save(path)
    print(f"geschrieben: {path}")


if __name__ == "__main__":
    build(sys.argv[1] if len(sys.argv) > 1 else "AVV_CertoClav_biotec.docx")
