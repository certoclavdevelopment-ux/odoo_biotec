#!/usr/bin/env python3
"""Kompaktes, kundenfähiges Protokoll des Discovery Calls.

Bewusst knapp: Der Termin diente dem Überblick, die belastbaren Details ergeben
sich erst aus der Auswertung der bereitgestellten Daten. Herzstück der Nacharbeit
ist die Datenanforderung (Dok.-Nr. DATA-2026-001).

Die ausführliche interne Fassung erzeugt protokoll_vorlage_discovery.py.

Aufruf:  python3 protokoll_kompakt.py <zieldatei.docx>
"""
import sys

from docx import Document
from docx.shared import Cm, Pt

from docx_bausteine import ACCENT, GREY, bullets, heading, note, seite_einrichten, table, titel


def absatz(doc, text, groesse=10.5, kursiv=False, abstand=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(abstand)
    r = p.add_run(text)
    r.font.size = Pt(groesse)
    r.font.italic = kursiv
    return p


def build(path):
    doc = Document()
    seite_einrichten(doc, "Protokoll Discovery Call  ·  Dok.-Nr. DISC-2026-001")
    titel(doc, "Discovery Call – Zusammenfassung",
          "Odoo für die biotec GmbH  ·  17. August 2026")

    table(doc, ["Feld", "Eintrag"], [4.2, 12.8], [
        ["Termin", "Montag, 17.08.2026, 12:30 – 13:30 Uhr, Microsoft Teams"],
        ["biotec GmbH", "Elbrachtsweg 76, 33332 Gütersloh · +49 5241 307200 · "
                        "info@biotec-gmbh.com"],
        ["Teilnehmer biotec", "Melanie Frank, Michael Brand, Nicole Krupa, "
                              "Dr. Andreas Bermpohl, Dr. Thomas Wilke, IT-Administration"],
        ["Teilnehmer Certania", "Moritz Gruber"],
        ["Teilnehmer CertoClav", "Michael Simon (Consultant & SPOC), Jonas Leitenmeier, "
                                 "Balázs Szaradics, Patrick Gottfried"],
        ["Protokoll", "Michael Simon"],
    ])

    heading(doc, "Zweck des Termins")
    absatz(doc,
           "Ziel war ein erster Überblick: Wie arbeitet biotec heute, welche Systeme sind im "
           "Einsatz, und wo kann Odoo helfen. Der Termin war ausdrücklich keine "
           "Detailerhebung – er sollte klären, ob und wie eine Zusammenarbeit sinnvoll ist.")

    heading(doc, "Was wir verstanden haben", kicker="Grober Überblick")
    note(doc, "Bewusst auf hoher Ebene. Die belastbaren Details ergeben sich aus den "
              "Unterlagen, die uns biotec bereitstellt.")
    bullets(doc, [
        "biotec arbeitet von zwei Standorten aus: Gütersloh als Hauptstandort mit Labor und "
        "Verwaltung, München mit einem Team für Hygienekontrollen.",
        "Der Kernprozess führt vom Auftrag über die Anlagenliste und die Probenahme vor Ort "
        "ins Labor, von dort über Bebrütung und Auszählung zum Gutachten und zur Rechnung.",
        "Ein zweiter Geschäftsbereich sind die bundesweiten VDI-Schulungen mit langer "
        "Vorplanung, Mindestteilnehmerzahl und internen sowie externen Dozenten.",
        "Heute im Einsatz: eine eigenentwickelte Anwendung mit mehreren Schnittstellen, ein "
        "separates Programm für die Gutachtenerstellung und OneDrive für die Ablage.",
        "Die Prüfer dokumentieren vor Ort mit dem Smartphone; Fotos gehören an den jeweiligen "
        "Datensatz.",
        "Lagerhaltung und Eigenproduktion sind überschaubar – Verbrauchsmaterial, einige "
        "Handelsartikel und wenige selbst hergestellte Chemikalien.",
    ])

    heading(doc, "Zwei Punkte, die wir festgehalten haben", kicker="Wichtig für die Umsetzung")
    bullets(doc, [
        "Die Dokumente – insbesondere Gutachten und Berichte – sollen für Ihre Kunden "
        "unverändert aussehen. Wir planen die Berichtsgestaltung entsprechend ein.",
        "Geschult wird am Ende des Projekts vor Ort bei Ihnen, gründlich und für einen "
        "festen Personenkreis. Eine reine Multiplikatoren-Schulung ist nicht gewünscht.",
    ])

    heading(doc, "Ansprechpartner")
    table(doc, ["Name", "E-Mail", "Themen"], [3.6, 5.4, 8.0], [
        ["Michael Brand", "michael.brand@biotec-gmbh.com",
         "Fachlich und technisch: Anlagen, Laborprozess, Gutachten, Verträge, Altsystem"],
        ["Nicole Krupa", "nicole.krupa@biotec-gmbh.com",
         "Kaufmännisch und organisatorisch: Kunden, Belege, Preise, Schulungen, Layout"],
        ["Michael Simon", "michael.simon@certoclav.com",
         "CertoClav: Beratung, Einrichtung, Projektsteuerung"],
    ])

    doc.add_page_break()

    heading(doc, "Der nächste Schritt: Datenanforderung", kicker="Herzstück der Nacharbeit")
    absatz(doc,
           "Wir haben Ihnen eine Datenanforderung zusammengestellt (Dok.-Nr. DATA-2026-001). "
           "Sie ist das eigentliche Arbeitspapier aus diesem Termin: 60 Positionen in zwölf "
           "Blöcken, jeweils mit Format und Zuständigkeit. Zwanzig Positionen sind als "
           "Paket 1 markiert – die brauchen wir zuerst.")
    bullets(doc, [
        "Eine Aufbereitung ist nicht nötig. Rohexporte sind uns lieber als schöne Tabellen – "
        "wir wollen sehen, wie die Daten wirklich aussehen.",
        "Unvollständig ist besser als nichts: Wir sagen Ihnen, was fehlt.",
        "Für den Upload haben wir einen OneDrive-Ordner freigegeben. Die Ordner dort sind so "
        "benannt wie die Blöcke der Datenanforderung.",
        "Personenbezogene Daten wie Teilnehmerlisten bitte anonymisiert oder mit Testnamen.",
    ])

    heading(doc, "Was danach passiert", kicker="Vom Upload zu Budget und Projektplan")
    absatz(doc,
           "Aus den bereitgestellten Unterlagen leiten wir die konkreten Anforderungen ab. "
           "Erst diese Auswertung erlaubt eine belastbare Aussage zu Aufwand, Budget und "
           "Zeitplan – eine Schätzung vor der Datenauswertung wäre geraten und für keine "
           "Seite hilfreich.")
    table(doc, ["Schritt", "Inhalt", "Wer"], [4.0, 9.6, 3.4], [
        ["1 · Upload", "biotec stellt die Unterlagen aus Paket 1 bereit", "biotec"],
        ["2 · Auswertung", "Wir prüfen Datenmenge, Datenqualität, Dokumentenvorlagen und die "
                           "Struktur des Altsystems und stellen Rückfragen", "CertoClav"],
        ["3 · Scoping", "Umfang je Phase, Reihenfolge der Themen, Abgrenzung dessen, was "
                        "bewusst später kommt", "CertoClav"],
        ["4 · Budget & Projektplan", "Aufwandsschätzung, Phasenplan mit Terminen, Vorschlag "
                                     "für den Go-live", "CertoClav"],
        ["5 · Entscheidung", "Freigabe und Kick-Off-Termin", "gemeinsam"],
    ])

    heading(doc, "Vertraulichkeit")
    absatz(doc,
           "Die Unterlagen, um die wir bitten, sind sensibel: Kundendaten, Gutachten, "
           "kaufmännische Zahlen und je nach Umfang auch die Datenbank und der Quellcode "
           "Ihrer eigenen Anwendung. Wir behandeln alles vertraulich, nutzen es "
           "ausschließlich für dieses Projekt und bewahren es getrennt von der übrigen "
           "Projektdokumentation auf.")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Auf Wunsch senden wir Ihnen gerne eine Vorlage für eine "
                  "Vertraulichkeitsvereinbarung (NDA) zu, die wir vor der Übergabe der "
                  "Unterlagen zeichnen. Sagen Sie uns einfach Bescheid.")
    r.font.size = Pt(10.5)
    r.font.bold = True
    absatz(doc,
           "Hinweis: Sobald personenbezogene Daten in größerem Umfang verarbeitet werden, "
           "kommt zusätzlich ein Vertrag zur Auftragsverarbeitung nach Art. 28 DSGVO in "
           "Betracht. Das stimmen wir gemeinsam ab.", groesse=9.5, kursiv=True)

    heading(doc, "Noch offen")
    note(doc, "Punkte, die im Termin nicht besprochen wurden und die wir im weiteren Verlauf "
              "klären.")
    bullets(doc, [
        "Wunschtermin für den Go-live und was ihn treibt.",
        "Budgetrahmen und Entscheidungsweg.",
        "Anzahl der späteren Odoo-Nutzer und deren Rollen.",
        "Gesellschaftsstruktur der beiden Standorte – daraus ergibt sich, ob Odoo als eine "
        "oder als mehrere Gesellschaften eingerichtet wird.",
        "Ob die Anforderungen der Certania-Gruppe an Beschaffung, Fakturierung und "
        "Projektbewertung Teil dieses Projekts sind.",
    ])

    heading(doc, "Rückfragen")
    p = doc.add_paragraph()
    r = p.add_run("Michael Simon  ·  michael.simon@certoclav.com\n"
                  "CertoClav Consulting  ·  CertoClav Sterilizer GmbH, Leonding")
    r.font.size = Pt(10)
    r.font.color.rgb = GREY

    doc.save(path)
    print(f"geschrieben: {path}")


if __name__ == "__main__":
    build(sys.argv[1] if len(sys.argv) > 1 else "Discovery_Call_Zusammenfassung.docx")
