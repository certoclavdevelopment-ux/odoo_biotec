#!/usr/bin/env python3
"""Erzeugt die Datenanforderung an biotec (kundenfähiges Word-Dokument).

Grundlage: Ergebnisse des Discovery Calls vom 17.08.2026. Die Liste ist nach
Themenblöcken sortiert; Paket 1 wird für das Scoping gebraucht, Paket 2 für die
Einrichtung, Paket 3 kann nachlaufen.

Aufruf:  python3 datenanforderung.py <zieldatei.docx>
"""
import sys

from docx import Document
from docx.shared import Cm, Pt

from docx_bausteine import GREY, bullets, heading, note, seite_einrichten, table, titel

# Verantwortliche – nach den im Discovery Call genannten Rollen
BRAND = "M. Brand"
FRANK = "M. Frank"
NKRUPA = "N. Krupa"
AKRUPA = "A. Krupa"
IT = "Westbomke"
GF = "Wilke / Bermpohl"

SPALTEN = ["Nr.", "Was wir brauchen", "Format", "Wer", "Paket", "Erledigt"]
BREITEN = [1.0, 7.4, 3.0, 2.2, 1.4, 1.8]

GELB = "FFF2A8"          # Hinterlegung der Paket-1-Zeilen
SPALTE_PAKET = 4         # Index der Paket-Spalte in den Zeilen unten


def absatz(doc, text, groesse=10.5, kursiv=False, abstand=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(abstand)
    r = p.add_run(text)
    r.font.size = Pt(groesse)
    r.font.italic = kursiv
    return p


def paket1_gelb(zeile):
    """Zeilen aus Paket 1 werden gelb hinterlegt."""
    return GELB if zeile[SPALTE_PAKET] == "1" else None

BLOECKE = [
    ("A · Kunden & Lieferanten", None, [
        ["A1", "Kundenliste: Kundennummer, Name, Adresse, USt-IdNr., Ansprechpartner, "
               "Telefon, E-Mail, Zahlungsbedingung, Preisgruppe", "Export CSV/XLSX", NKRUPA, "1", "☐"],
        ["A2", "Abweichende Liefer- und Rechnungsadressen; Zuordnung bei Verbünden "
               "(z. B. Klinikverbund mit mehreren Häusern)", "CSV/XLSX", NKRUPA, "2", "☐"],
        ["A3", "Lieferantenliste: Lieferantennummer, Name, Adresse, USt-IdNr., "
               "Zahlungsbedingung – inklusive Fremdlabore", "Export CSV/XLSX", AKRUPA, "1", "☐"],
    ]),

    ("B · Anlagen & Objekte", "Das zentrale Objekt Ihres Geschäfts – wichtigster Datenblock.", [
        ["B1", "Anlagen-/Objektstamm: Anlagennummer, Kunde, Gebäude/Standort, Anlagentyp, "
               "Baujahr, Prüfnorm (VDI 6022 / 2047), Prüfintervall, letzte und nächste Prüfung",
         "Export CSV/XLSX", BRAND, "1", "☐"],
        ["B2", "Zwei bis drei Beispiel-Anlagenlisten so, wie sie heute erzeugt werden – "
               "einmal als Datei und einmal als Papierausdruck (Scan/PDF)", "PDF + Datei", BRAND, "1", "☐"],
        ["B3", "Prüfpläne und Checklisten je Norm und Anlagentyp", "PDF/Word", FRANK, "2", "☐"],
        ["B4", "Turnusregeln: welches Intervall gilt für welchen Anlagentyp, wie werden "
               "Folgeaufträge heute ausgelöst", "kurze Beschreibung", BRAND, "2", "☐"],
    ]),

    ("C · Labor", None, [
        ["C1", "Probenbegleitschein und Laborjournal als Muster", "PDF", FRANK, "2", "☐"],
        ["C2", "Analysemethoden: Medien, Bebrütungszeiten und -temperaturen, Grenzwerte, "
               "Bewertungsschema", "Übersicht", FRANK, "2", "☐"],
        ["C3", "Kennzeichnung von Proben und Platten: Nummernlogik, Etiketten, Barcodes – "
               "falls vorhanden", "Beschreibung/Muster", FRANK, "2", "☐"],
        ["C4", "Fremdlabor: Beispiel einer Beauftragung und des zurückgelieferten Ergebnisberichts",
         "PDF", FRANK, "2", "☐"],
    ]),

    ("D · Gutachten, Belege & Layout", "Kritisch: Die Dokumente sollen exakt so aussehen wie "
                                       "heute. Dafür brauchen wir Muster und Vorlagendateien.", [
        ["D1", "Je Berichtstyp ein fertiges Gutachten als PDF – anonymisiert genügt",
         "PDF", BRAND, "1", "☐"],
        ["D2", "Die Vorlagendateien des Gutachtenprogramms (Report-Templates, Word-Vorlagen) – "
               "nicht nur das Ergebnis", "Originaldateien", IT, "1", "☐"],
        ["D3", "Angebot, Auftragsbestätigung, Lieferschein, Rechnung, Mahnung – je ein Muster",
         "PDF", NKRUPA, "1", "☐"],
        ["D4", "Briefpapier, Logo als Vektordatei, Hausfarben, verwendete Schriftarten",
         "AI/SVG/EPS + PDF", NKRUPA, "1", "☐"],
        ["D5", "Standardtexte und Textbausteine: Bewertungsformulierungen, Angebotstexte, "
               "E-Mail-Signaturen, AGB", "Word/PDF", NKRUPA, "2", "☐"],
    ]),

    ("E · Verkauf & Einkauf", None, [
        ["E1", "Angebote und Auftragsbestätigungen der letzten Wochen", "PDF", NKRUPA, "1", "☐"],
        ["E2", "Ausgangsrechnungen der letzten Wochen", "PDF", NKRUPA, "1", "☐"],
        ["E3", "Preisliste(n), Rabattlogik, kundenspezifische Sonderpreise",
         "XLSX", NKRUPA, "2", "☐"],
        ["E4", "Verträge: Rahmen-, Wartungs- und Prüfverträge als Beispiele",
         "PDF", BRAND, "1", "☐"],
        ["E5", "Bestellhistorie und offene Bestellungen", "XLSX", AKRUPA, "2", "☐"],
        ["E6", "Reklamationsliste und der heutige Umgang damit", "XLSX + Beschreibung", BRAND, "2", "☐"],
    ]),

    ("F · Artikel & Lager", None, [
        ["F1", "Artikelliste mit Bestand: Artikelnummer, Bezeichnung, Einheit, Lagerort, "
               "Mindestbestand, Lieferant, Einkaufs- und Verkaufspreis", "Export XLSX", AKRUPA, "1", "☐"],
        ["F2", "Chargen- und Verfallsdatenpflicht bei Chemikalien und Nährmedien – ja oder nein, "
               "und wie wird das heute dokumentiert", "kurze Antwort", FRANK, "2", "☐"],
        ["F3", "Eigenproduktion: Rezepturen der selbst hergestellten Chemikalien, "
               "typische Chargengrößen", "Beschreibung", FRANK, "2", "☐"],
    ]),

    ("G · Schulungen & Veranstaltungen", "Im Discovery Call als eigener Geschäftsbereich "
                                         "beschrieben – bisher nicht in der Datenliste enthalten.", [
        ["G1", "Kurskatalog: Kursart (VDI 6022 Kat. A/B, VDI 2047), Dauer, Preis, "
               "Mindest- und Maximalteilnehmer", "XLSX", NKRUPA, "1", "☐"],
        ["G2", "Seminarkalender: Termine, Orte, Räume, Vorlaufzeiten der Raumbuchung",
         "XLSX", NKRUPA, "2", "☐"],
        ["G3", "Anmeldeweg heute: Formular, Website, E-Mail? Bestätigungs- und Erinnerungsmails",
         "Screenshots/Muster", NKRUPA, "2", "☐"],
        ["G4", "Teilnehmerliste – als Muster genügen wenige Beispielzeilen, anonymisiert. "
               "Sollen historische Teilnehmer für Bescheinigungen übernommen werden, brauchen "
               "wir sie später mit echten Angaben",
         "XLSX", NKRUPA, "2", "☐"],
        ["G5", "Teilnahmebescheinigung und Zertifikatsvorlage", "Originaldatei + PDF", NKRUPA, "2", "☐"],
        ["G6", "Dozentenliste (intern/extern), Honorarsätze der externen Dozenten",
         "XLSX", BRAND, "2", "☐"],
        ["G7", "Inhouse-Schulungen: Angebotsmuster und Preislogik", "PDF", NKRUPA, "2", "☐"],
    ]),

    ("H · Buchhaltung & Finanzen", None, [
        ["H1", "Vollständiger Kontenrahmen mit Kontenbeschriftung – nicht nur Ertrag und "
               "Aufwand, und mit Angabe des Standards (SKR03 / SKR04 / eigener)",
         "Export CSV/XLSX", AKRUPA, "1", "☐"],
        ["H2", "Steuerschlüssel und Steuersätze; Sonderfälle wie Reverse Charge, "
               "Auslandsleistungen, steuerfreie Umsätze", "Übersicht", AKRUPA, "2", "☐"],
        ["H3", "Zahlungsarten und Zahlungsbedingungen; Bankkonten; SEPA-Lastschriftmandate",
         "Übersicht", AKRUPA, "2", "☐"],
        ["H4", "Mahnstufen und Mahntexte", "Word/PDF", AKRUPA, "2", "☐"],
        ["H5", "Offene Posten Debitoren und Kreditoren zum Stichtag", "XLSX", AKRUPA, "2", "☐"],
        ["H6", "Saldenliste bzw. Eröffnungsbilanz zum geplanten Umstellungsstichtag",
         "PDF/XLSX", AKRUPA, "2", "☐"],
        ["H7", "DATEV: Kanzlei, Berater- und Mandantennummer, heutiges Exportformat",
         "Angabe", AKRUPA, "2", "☐"],
        ["H8", "Kostenstellen- und Kostenträgerstruktur, falls verwendet", "Übersicht", AKRUPA, "2", "☐"],
    ]),

    ("I · Altsystem (Delphi-Anwendung)", "Für die Migration ist die Datenbank wichtiger als der "
                                         "Quellcode – bitte beides, aber die Datenbank zuerst.", [
        ["I1", "Datenbanksystem und Version (z. B. Firebird, MS SQL) sowie das Schema: "
               "Tabellen, Felder, Beziehungen", "Schema-Export/PDF", IT, "1", "☐"],
        ["I2", "Datenbank-Sicherung oder Vollexport aller Tabellen", "Dump oder CSV je Tabelle", IT, "1", "☐"],
        ["I3", "Quellcode der hauseigenen Software inklusive Delphi-Version und "
               "verwendeter Komponenten", "Repository/Archiv", IT, "2", "☐"],
        ["I4", "Schnittstellenliste: welches Gegensystem, welche Richtung, welches Format, "
               "welche Frequenz", "Übersicht", IT, "1", "☐"],
        ["I5", "Screenshots der Hauptmasken oder eine kurze Bildschirmaufnahme eines "
               "kompletten Durchlaufs", "PNG/Video", BRAND, "1", "☐"],
        ["I6", "Zugang zu einem Test- oder Kopiesystem für Rückfragen", "Zugang", IT, "2", "☐"],
        ["I7", "Lizenz- und Wartungssituation, betreuender Dienstleister", "Angabe", IT, "2", "☐"],
    ]),

    ("J · IT & Infrastruktur", None, [
        ["J1", "Systemübersicht: Server, Betriebssysteme, Standorte, Backup-Verfahren",
         "Übersicht", IT, "2", "☐"],
        ["J2", "Microsoft 365: Tenant, Struktur von OneDrive und SharePoint, Postfächer",
         "Übersicht", IT, "2", "☐"],
        ["J3", "Endgeräte: Anzahl der Pixel-8-Geräte, weitere Geräte, Netzabdeckung in "
               "Technikzentralen (Offline-Fähigkeit nötig?)", "Angabe", IT, "1", "☐"],
        ["J4", "Nutzerliste mit Rolle und Standort – Grundlage für Odoo-Nutzer und Rechte",
         "XLSX", GF, "2", "☐"],
        ["J5", "Datenschutz: Vorlage für den Auftragsverarbeitungsvertrag, Löschfristen, "
               "zulässige Speicherorte", "PDF", GF, "2", "☐"],
    ]),

    ("K · Organisation & Struktur", None, [
        ["K1", "Handelsregisterauszüge biotec und biotec Süd sowie die Gesellschaftsstruktur – "
               "entscheidet über die Mehrgesellschafts-Einrichtung in Odoo", "PDF", GF, "1", "☐"],
        ["K2", "Organigramm mit Standorten und Teams", "PDF", GF, "2", "☐"],
        ["K3", "Zeiterfassung: wird heute erfasst, womit? Interne Stunden- und Kostensätze",
         "Angabe", GF, "2", "☐"],
        ["K4", "Reisekosten- und Spesenprozess, Fahrzeuge", "Beschreibung", AKRUPA, "3", "☐"],
        ["K5", "Archivierungs- und Aufbewahrungsvorgaben (GoBD, Nachweispflichten aus VDI)",
         "Beschreibung", GF, "3", "☐"],
    ]),

    ("L · Certania-Gruppe", "Nur relevant, falls die Gruppenanforderungen zum Projekt gehören – "
                            "das ist noch offen.", [
        ["L1", "Vorlage des Reporting Package und die Melderhythmen", "XLSX/PDF", GF, "2", "☐"],
        ["L2", "Konzernkontenplan und Mapping auf die lokalen Konten, falls vorhanden",
         "XLSX", AKRUPA, "3", "☐"],
        ["L3", "Intercompany-Leistungen zwischen Gruppengesellschaften", "Übersicht", AKRUPA, "3", "☐"],
    ]),
]


def build(path):
    doc = Document()
    seite_einrichten(doc, "Datenanforderung  ·  Dok.-Nr. DATA-2026-001")
    titel(doc, "Datenanforderung",
          "Odoo-Einführung biotec GmbH  ·  Was wir für Scoping und Einrichtung brauchen")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.add_run(
        "Bitte stellen Sie die folgenden Unterlagen bereit. "
        "Eine Aufbereitung ist nicht nötig – Rohexporte und Beispieldateien genügen, "
        "auch wenn sie unvollständig oder unsauber sind. Lieber ein unfertiger Export "
        "als keiner: Wir sehen daraus, wie Ihre Daten wirklich aussehen.")

    heading(doc, "Ihre Ansprechpartner für uns", kicker="Koordination")
    table(doc, ["Name", "E-Mail", "Rolle", "Zuständig für"], [3.2, 4.8, 4.0, 5.0], [
        ["Michael Brand", "michael.brand@biotec-gmbh.com",
         "Projektleiter Hygieneinspektion, Begleiter ERP-Projekt",
         "Fachliche und technische Themen: Anlagen, Laborprozess, Gutachten, Verträge, "
         "Altsystem"],
        ["Nicole Krupa", "nicole.krupa@biotec-gmbh.com",
         "Rechnungen, Schulungen, Werbung",
         "Kaufmännische und organisatorische Themen: Kunden, Belege, Preise, Schulungen, "
         "Layout"],
    ])
    note(doc, "Beide bündeln die Rückmeldungen auf Ihrer Seite und stimmen sich mit den "
              "übrigen Beteiligten ab. Wir wenden uns mit Rückfragen an sie, statt einzelne "
              "Personen direkt anzusprechen.")

    heading(doc, "Reihenfolge", kicker="Drei Pakete")
    bullets(doc, [
        "Paket 1 – in den Tabellen gelb hinterlegt. Das brauchen wir zuerst, für Scoping "
        "und Aufwandsschätzung. Auch unvollständig.",
        "Paket 2 – für die Einrichtung. Kann nachlaufen.",
        "Paket 3 – für den Feinschliff vor dem Go-live.",
    ])
    table(doc, ["Legende"], [17.0],
          [["Gelb hinterlegte Zeilen = Paket 1. Diese Positionen brauchen wir zuerst."]],
          zeilen_fuellung=lambda z: GELB)

    heading(doc, "So liefern Sie am besten", kicker="Format & Weg")
    bullets(doc, [
        "Exporte bevorzugt als CSV oder Excel, direkt aus dem System – ohne Nachbearbeitung.",
        "Dokumentenmuster als PDF, Vorlagen zusätzlich als Originaldatei.",
        "Bereitstellung über den geteilten Projektordner, nicht per E-Mail-Anhang. "
        "Bei großen Datenbanksicherungen sprechen Sie uns an, wir richten einen Upload ein.",
        "Personenbezogene Daten: siehe den eigenen Abschnitt unten – Stammdaten brauchen wir "
        "mit echten Angaben, bei reinen Mustern genügen Beispielzeilen.",
        "Ein Ordner je Block (A bis L) hilft uns bei der Zuordnung.",
    ])

    heading(doc, "Personenbezogene Daten", kicker="Was echt sein muss und was nicht")
    absatz(doc,
           "Eine vollständige Anonymisierung ist bei den meisten Positionen weder möglich noch "
           "sinnvoll: Kunden, Ansprechpartner und Anlagen sollen ja genau so in Odoo stehen. "
           "Deshalb unterscheiden wir zwei Fälle.")
    table(doc, ["Fall", "Beispiele", "Was wir brauchen"], [3.4, 6.6, 7.0], [
        ["Wird nach Odoo übernommen",
         "A1 Kunden mit Ansprechpartnern, A3 Lieferanten, B1 Anlagen, E1/E2 Angebote und "
         "Rechnungen, F1 Artikel, H5 offene Posten, J4 Nutzerliste",
         "**Echte Daten.** Diese Datensätze sind das Migrationsziel – anonymisiert wären sie "
         "wertlos. Rechtliche Grundlage: diese Vertraulichkeitsvereinbarung, bei Bedarf "
         "ergänzt um einen Vertrag zur Auftragsverarbeitung nach Art. 28 DSGVO"],
        ["Nur Struktur und Layout",
         "C1 Probenbegleitschein, D1 Gutachten als Muster, D3 Belegmuster, G4 Teilnehmerliste",
         "**Beispielzeilen genügen.** Personennamen dürfen geschwärzt, ersetzt oder erfunden "
         "sein – wir brauchen nur Felder, Aufbau und Layout, nicht die Inhalte"],
    ])
    absatz(doc,
           "Wenn Sie unsicher sind, ob eine Position in den ersten oder zweiten Fall gehört: "
           "fragen Sie kurz nach. Wir schauen gemeinsam auf die Datei, das ist in zwei Minuten "
           "geklärt.", groesse=9.5, kursiv=True)

    heading(doc, "Wer liefert was", kicker="Zuordnung in der Tabelle")
    note(doc, "In der Spalte „Wer“ steht, wer die Unterlage am schnellsten zur Hand hat – "
              "ein Vorschlag auf Basis der im Discovery Call genannten Funktionen. "
              "Die Koordination läuft über Herrn Brand und Frau Krupa; bitte intern gern "
              "anders verteilen, wenn es so besser passt.")
    table(doc, ["Kürzel", "Person", "Blöcke"], [2.4, 5.6, 9.0], [
        [BRAND, "Michael Brand – Ansprechpartner (fachlich/technisch)",
         "Anlagen, Gutachten, Verträge, Reklamationen, Altsystem-Masken"],
        [NKRUPA, "Nicole Krupa – Ansprechpartnerin (kaufmännisch/organisatorisch)",
         "Kunden, Belege, Preise, Schulungen, Layout"],
        [FRANK, "Melanie Frank – Laborkoordination, übergeordnete Prozesse",
         "Labor, Analysemethoden, Prüfpläne, Eigenproduktion"],
        [AKRUPA, "Annette Krupa – Buchhaltung",
         "Kontenrahmen, Steuern, Zahlungen, offene Posten, Artikel"],
        [IT, "IT-Administration (Westbomke)",
         "Datenbank, Schnittstellen, Quellcode, Infrastruktur"],
        [GF, "Dr. Wilke / Dr. Bermpohl",
         "Gesellschaftsstruktur, Organigramm, Nutzerliste, Datenschutz"],
    ])

    doc.add_page_break()

    for name, hinweis, zeilen in BLOECKE:
        heading(doc, name)
        if hinweis:
            note(doc, hinweis)
        table(doc, SPALTEN, BREITEN, zeilen, zeilen_fuellung=paket1_gelb)

    heading(doc, "Vertraulichkeit")
    p = doc.add_paragraph()
    r = p.add_run(
        "Quellcode, Datenbanksicherungen und Kundendaten behandeln wir vertraulich und "
        "verwenden sie ausschließlich für dieses Projekt. Sie werden getrennt von der "
        "Projektdokumentation aufbewahrt und nach Projektende auf Wunsch gelöscht. "
        "Falls für die Übergabe eine Vertraulichkeitsvereinbarung nötig ist, sagen Sie "
        "uns Bescheid – wir zeichnen sie vor der Übergabe.")
    r.font.size = Pt(10)

    heading(doc, "Rückfragen")
    p = doc.add_paragraph()
    r = p.add_run("Michael Simon  ·  michael.simon@certoclav.com  ·  CertoClav Consulting\n"
                  "Auf Ihrer Seite: Michael Brand (michael.brand@biotec-gmbh.com) und "
                  "Nicole Krupa (nicole.krupa@biotec-gmbh.com)")
    r.font.size = Pt(10)
    r.font.color.rgb = GREY

    doc.save(path)
    print(f"geschrieben: {path}")
    gesamt = sum(len(z) for _, _, z in BLOECKE)
    p1 = sum(1 for _, _, z in BLOECKE for r in z if r[4] == "1")
    print(f"{gesamt} Positionen in {len(BLOECKE)} Blöcken, davon {p1} in Paket 1")


if __name__ == "__main__":
    build(sys.argv[1] if len(sys.argv) > 1 else "Datenanforderung_biotec.docx")
