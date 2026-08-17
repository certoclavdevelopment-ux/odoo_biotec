#!/usr/bin/env python3
"""Erzeugt Vorlage und ausgefüllte Protokolle für den Discovery Call (biotec GmbH).

Aufbau folgt der Discovery-Call-Präsentation
(03_Praesentationen/2026-08-17_Discovery_Call/Discovery_Call_Biotec_CertoClav.pdf).

Aufruf:
  leere Vorlage:  python3 protokoll_vorlage_discovery.py <zieldatei.docx>
  Termin gefüllt: python3 protokoll_vorlage_discovery.py <zieldatei.docx> 2026-08-17

Inhalte je Termin stehen in TERMINE. Fragen ohne Antwort bekommen Schreiblinien,
Fragen mit Antwort bekommen die Antwort als Aufzählung.
"""
import sys

from docx import Document
from docx.shared import Cm, Pt

from docx_bausteine import (OFFEN, answer_lines, bullets, heading, note, question,
                           seite_einrichten, table, titel)


# --------------------------------------------------------------------------- Fragenkatalog
# (Schlüssel, Frage, Hinweis) – Schlüssel verbindet Frage und Antwort je Termin.
F_SCHMERZ = [
    ("warum_jetzt", "Warum sprechen wir gerade jetzt über ein ERP?", None),
    ("handarbeit", "Was funktioniert heute nicht – oder nur mit viel Handarbeit?",
     "(Doppelerfassung, Zettelwirtschaft, Excel-Inseln)"),
    ("zeitverlust", "Wo geht im Alltag am meisten Zeit verloren?",
     "(Einsatzplanung, Berichte, Rechnungsstellung, Kursorganisation)"),
    ("ausloeser", "Was ist der Auslöser, das jetzt anzugehen?",
     "(Wachstum, Personalwechsel, Kundenanforderungen, Altsystem)"),
    ("erfolg_12m", "Woran würden Sie in 12 Monaten festmachen, dass sich das Projekt gelohnt hat?", None),
    ("no_go", "Was soll auf keinen Fall schlechter werden als heute?", None),
]

F_PROZESSE = [
    ("software_heute", "Welche Software ist heute im Einsatz – und wofür?",
     "(Buchhaltung/DATEV, Planung, Labor, Kursbuchung, Office)"),
    ("auftragsdurchlauf", "Wie läuft ein Auftrag heute durch – von der Anfrage bis zur Rechnung?",
     "(Inspektion, Beratung, Schulung im Vergleich)"),
    ("daten_qualitaet", "Welche Daten liegen vor – und in welcher Qualität?",
     "(Kunden, Anlagen/Objekte, Prüfberichte, Vorlagen, Kurskatalog)"),
    ("vor_ort", "Wie arbeiten die Prüfer vor Ort?",
     "(mobil/offline, Checklisten, Fotos, Unterschriften, Berichtserstellung)"),
    ("bleiben_ersetzen",
     "Welche Systeme müssen bleiben und angebunden werden – was darf Odoo ersetzen?", None),
]

F_FINANCE = [
    ("scope_poc", "Gehören die Gruppenanforderungen (P2P/O2C, POC/WIP) zum biotec-Projekt – oder separat?", None),
    ("reporting", "Welche Anforderungen bestehen an das Reporting Package der Gruppe?",
     "(Inhalte, Frequenz, Empfänger, Format)"),
    ("projektbewertung", "Laufen mehrjährige/mehrmonatige Projekte, die nach Fertigstellungsgrad bewertet werden?",
     "(Festpreis, Meilensteinabrechnung, Fremdlabor)"),
    ("bilanzierung", "Bilanzierung nach HGB, IFRS oder beidem? Wer ist Steuerberater / Wirtschaftsprüfer?", None),
]

F_TEAM = [
    ("spoc", "Wer wäre Ihr Hauptansprechpartner für die Iterationen?",
     "(stimmt Feedbackrunden ab, schult später intern)"),
    ("kapazitaet", "Wie viel Zeit kann diese Person pro Woche realistisch einbringen?",
     "(Richtwert: wenige Stunden für Demos & Freigaben)"),
    ("nutzerzahl", "Wie viele Mitarbeiter würden mit Odoo arbeiten?",
     "(Innendienst, Prüfer, Labor, Verwaltung)"),
    ("golive", "Gibt es einen Wunschtermin für den Go-live – und was treibt ihn?", None),
    ("budget", "In welchem Budgetrahmen bewegen wir uns – und wer entscheidet?",
     "(Basis für Scoping & Phasierung, keine Festlegung heute)"),
]

PROFIL_LEER = [
    ["Bundesweites Hygieneinstitut, VDI-Schulungspartner (VDI 6022 / 2047)", "☐ ja ☐ nein", ""],
    ["Inspektionen vor Ort: RLT-Anlagen, Verdunstungskühlanlagen, Luftentkeimung", "☐ ja ☐ nein", ""],
    ["Beratung & Gefährdungsbeurteilungen (Trinkwasser-, Krankenhaushygiene)", "☐ ja ☐ nein", ""],
    ["Bundesweite VDI-Schulungen, Termine online buchbar", "☐ ja ☐ nein", ""],
    ["Labor & Forschung, Entwicklung neuer Analyseverfahren", "☐ ja ☐ nein", ""],
    ["Standort Gütersloh (Elbrachtsweg 76, 33332)", "☐ ja ☐ nein", ""],
]

APPS_LEER = [
    ["CRM & Verkauf", "Anfragen, Angebote & Aufträge", "☐ ja ☐ nein ☐ offen", ""],
    ["Rechnungsstellung", "Rechnungen & Zahlungen", "☐ ja ☐ nein ☐ offen", ""],
    ["Einkauf & Lager", "Labor- & Prüfmaterial", "☐ ja ☐ nein ☐ offen", ""],
    ["Außendienst", "Einsatzplanung & Prüfberichte vor Ort", "☐ ja ☐ nein ☐ offen", ""],
    ["Projekte & Zeiterfassung", "Beratungsaufträge & Gutachten", "☐ ja ☐ nein ☐ offen", ""],
    ["Veranstaltungen & E-Learning", "VDI-Schulungen inkl. Online-Buchung", "☐ ja ☐ nein ☐ offen", ""],
    ["Abonnements", "Wiederkehrende Prüfintervalle & Verträge", "☐ ja ☐ nein ☐ offen", ""],
    ["Qualität & Dokumente", "Proben, Berichte & Nachweise", "☐ ja ☐ nein ☐ offen", ""],
    ["Buchhaltung", "Finanzbuchhaltung, DATEV-Anbindung", "☐ ja ☐ nein ☐ offen", ""],
    ["Produktion", "bewusst nicht eingeplant – klären", "☐ ja ☐ nein ☐ offen", ""],
    ["Website / E-Commerce", "bewusst nicht eingeplant – klären", "☐ ja ☐ nein ☐ offen", ""],
    ["Weitere:", "", "☐ ja ☐ nein ☐ offen", ""],
]

AUFGABEN_LEER = [
    ["1", "Zusammenfassung / Protokoll des Discovery Calls versenden", "CertoClav", "", "☐ offen"],
    ["2", "Beispielhafte Rohdaten & Vorlagen teilen (keine Aufbereitung nötig)", "biotec", "", "☐ offen"],
    ["3", "Scoping & Angebot: Phasierung, Aufwandsschätzung, Budgetbasis", "CertoClav", "", "☐ offen"],
    ["4", "Entscheidung & Kick-Off: Hauptansprechpartner benennen, Termin", "Gemeinsam", "", "☐ offen"],
    ["5", "", "", "", "☐ offen"],
    ["6", "", "", "", "☐ offen"],
    ["7", "", "", "", "☐ offen"],
]

LEERE_META = [
    ["Datum / Uhrzeit", ""],
    ["Dauer (geplant 45–60 Min.)", ""],
    ["Ort / Tool", ""],
    ["Teilnehmer biotec", ""],
    ["Teilnehmer CertoClav", "Michael Simon (Consultant & SPOC)"],
    ["Weitere Teilnehmer (Certania)", ""],
    ["Protokoll", "Michael Simon"],
    ["Aufzeichnung", "☐ ja   ☐ nein   (Einverständnis eingeholt: ☐ ja  ☐ nein)"],
]


# --------------------------------------------------------------------------- Termindaten
TERMINE = {
    "2026-08-17": {
        "meta": [
            ["Datum / Uhrzeit", "Montag, 17.08.2026, 12:30 – 13:30 Uhr"],
            ["Dauer", "60 Minuten"],
            ["Ort / Tool", "Microsoft Teams (Videokonferenz, via Calendly gebucht)"],
            ["Betreff der Einladung", "Discovery Call biotec GmbH"],
            ["Organisator", "Michael Simon, michael.simon@certoclav.com"],
            ["Protokoll", "Michael Simon"],
            ["Rückmeldungen (Stand Einladung)", "4 Zusagen, 0 vorläufig, 0 Absagen"],
            ["Aufzeichnung", "☐ ja   ☐ nein   (Einverständnis eingeholt: ☐ ja  ☐ nein)"],
        ],
        "teilnehmer": [
            ["Michael Simon", "michael.simon@certoclav.com", "CertoClav",
             "Consultant & SPOC, Organisator", "Organisator"],
            ["Dr. Thomas Wilke", "twilke@certania.com", "Certania / biotec",
             "Direktor Food (Certania), CEO biotec", "keine Antwort"],
            ["Dr. Andreas Bermpohl", "abermpohl@t-online.de", "biotec (extern)",
             "Ansprechpartner Richtung biotec, Prokura", "keine Antwort"],
            ["Melanie Frank", "melanie.frank@biotec-gmbh.com", "biotec",
             "Promovierte Biologin; koordiniert die Labortätigkeiten und übergeordnete "
             "Prozesse; Hygieneinspektion", "keine Antwort"],
            ["Michael Brand", "michael.brand@biotec-gmbh.com", "biotec",
             "Projektleiter Hygieneinspektion seit 2020; Organisation der Kunden und "
             "Übergabe der Gutachten; Begleiter des ERP-Projekts", "keine Antwort"],
            ["Nicole Krupa", "nicole.krupa@biotec-gmbh.com", "biotec",
             "seit einem Jahr bei biotec; enger Kontakt zu den Projektleitern; "
             "Rechnungen, Schulungen, Werbung", "keine Antwort"],
            ["Annette Krupa", "–", "biotec", "Buchhaltung", "nicht eingeladen"],
            ["unbekannt (mw@)", "mw@westbomke.com", "Westbomke (?)",
             "IT-Administration", "keine Antwort"],
            ["Moritz Gruber", "mgruber@certania.com", "Certania (optional)",
             "CEO / Managing Shareholder, Initiator", "keine Antwort"],
            ["Jonas Leitenmeier", "–", "CertoClav", "AI-Ops-Lead", "✓ zugesagt"],
            ["Balázs Szaradics", "–", "CertoClav", "AI-Operator", "✓ zugesagt"],
            ["Patrick Gottfried", "–", "CertoClav", "AI-Operator", "✓ zugesagt"],
        ],
        "hinweis": ("Annette Krupa (Buchhaltung) war nicht eingeladen, ist für das ERP-Projekt "
                    "aber relevant. Bei Nicole Krupa steht in der Mitschrift „Biotec Stetic“ – "
                    "Zuordnung zu klären. Anwesenheit im Termin abweichend von der Einladung "
                    "bitte korrigieren."),

        "kernaussagen": [
            "biotec besteht aus drei Standorten – Gütersloh als Hauptstandort, München "
            "(Hygienekontrollen) und dem Technologiepark Mittweida (Analytik Boden und "
            "Wasser). Gesucht wird ein ERP für eine kleine Gesellschaft.",
            "Das Altsystem ist eine Delphi-Applikation mit vielen Schnittstellen; die Berichte "
            "erzeugt ein separates Gutachtenprogramm, die Ablage läuft über OneDrive.",
            "Harte Anforderung: Die Dokumente – insbesondere Gutachten und Berichte – müssen "
            "exakt so aussehen wie bisher.",
            "Das Deck-Modell Train-the-Trainer wurde abgelehnt. Gewünscht ist eine gründliche "
            "Vor-Ort-Schulung am Projektende für drei bis vier namentlich benannte Personen.",
            "Zwei Geschäftsbereiche prägen den Scope: der Laborprozess "
            "(Probenahme → Bebrütung → Auszählung → Gutachten → Rechnung) und das "
            "Schulungsgeschäft (bundesweite VDI-Seminare mit langer Vorplanung).",
        ],

        "antworten": {
            # 1 · Arbeitsmodell
            "modell_reaktion": [
                "Train-the-Trainer wird abgelehnt. Stattdessen: gründliche Schulung am Ende des "
                "Projekts, vor Ort beim Kunden.",
                "Zu schulen sind Stefan, Michael Brand und Markus sowie eine weitere Person. "
                "Die Notizen nennen sowohl „3 Personen“ als auch „4 Leute“ – Anzahl und Namen "
                "sind zu bestätigen.",
                "Konsequenz für das Scoping: Der Schulungsaufwand liegt bei CertoClav statt beim "
                "Hauptansprechpartner. Das Arbeitsmodell aus dem Deck ist an dieser Stelle "
                "anzupassen.",
            ],
            # 3 · Schmerzpunkte & Ziele
            "warum_jetzt": OFFEN,
            "handarbeit": [
                "Die Anlagenliste wird in einer kleinen Zusatzsoftware erstellt und auf Papier "
                "ausgedruckt; im Labor wird nach diesem Papierdokument gearbeitet.",
                "Ergebnisse (Auszählung, Fotos) werden anschließend manuell zurückgeführt und am "
                "Datensatz hinterlegt.",
            ],
            "zeitverlust": OFFEN,
            "ausloeser": OFFEN,
            "erfolg_12m": OFFEN,
            "no_go": [
                "Die Dokumente müssen exakt so aussehen wie bisher – Gutachten und Berichte sind "
                "gegenüber den Kunden gesetzt.",
            ],
            # 4 · Prozesse, Systeme & Daten
            "software_heute": [
                "Altsystem: Delphi-Applikation mit vielen Schnittstellen.",
                "Separates Gutachtenprogramm erzeugt die Berichte.",
                "Kleine Zusatzsoftware für die Erstellung der Anlagenliste.",
                "Dokumentenablage in OneDrive.",
                "Buchhaltungssystem und Kursverwaltung: im Termin nicht erhoben.",
            ],
            "auftragsdurchlauf": [
                "Auftrag geht ein.",
                "Anlagenliste wird in der kleinen Software erstellt und in Papierform ausgedruckt.",
                "Im Labor wird gemäß diesem Dokument gearbeitet; Proben werden genommen.",
                "Proben werden nach Gütersloh gebracht.",
                "Bebrütung; Auszählung nach 3 oder 5 Tagen.",
                "Auszählung, Fotos, Bild am Datensatz hinterlegen.",
                "Gutachtenprogramm erstellt den Bericht; Ablage in OneDrive.",
                "Nach Sichtung durch das Labor werden die Proben autoklaviert und entsorgt.",
                "Ist das Gutachten fertig, wird eine Info erzeugt und die Rechnung gestellt.",
            ],
            "daten_qualitaet": [
                "Im Termin nicht systematisch erhoben.",
                "Anzufordern: Anlagenlisten, Gutachten- und Berichtsvorlagen, Kurskatalog, "
                "Artikel- und Kundenstamm aus der Delphi-Applikation.",
            ],
            "vor_ort": [
                "Fotos müssen direkt vor Ort per Smartphone möglich sein – im Einsatz sind "
                "Google Pixel 8.",
                "Fotos werden am Datensatz bzw. an der Anlage hinterlegt.",
            ],
            "bleiben_ersetzen": [
                "Ablösekandidaten: Delphi-Applikation, Gutachtenprogramm, Zusatzsoftware für "
                "die Anlagenliste.",
                "Bindende Randbedingung: Layout der erzeugten Dokumente bleibt unverändert.",
                "Offen: Welche Schnittstellen die Delphi-Applikation bedient und ob OneDrive als "
                "Ablage bestehen bleibt.",
            ],
            # 6 · Finance / Certania
            "scope_poc": OFFEN,
            "reporting": OFFEN,
            "projektbewertung": OFFEN,
            "bilanzierung": OFFEN,
            # 7 · Team, Zeitrahmen & Budget
            "spoc": [
                "Nicht formell benannt. Nach Rollenlage ist Michael Brand der naheliegende "
                "Kandidat: Projektleiter Hygieneinspektion und ausdrücklich „Begleiter des "
                "ERP-Systems“.",
                "Namentlich für die Vor-Ort-Schulung vorgesehen: Stefan, Michael Brand, Markus "
                "(+ eine weitere Person).",
                "Fachliche Breite für die Validierung: Melanie Frank (Labor und übergeordnete "
                "Prozesse), Nicole Krupa (Rechnungen, Schulungen, Werbung), Annette Krupa "
                "(Buchhaltung), mw@westbomke.com (IT-Administration).",
            ],
            "kapazitaet": OFFEN,
            "nutzerzahl": [
                "Grobe Einordnung aus dem Termin: „ERP für eine kleine Gesellschaft“.",
                "Bekannte Teams: München 2–3 Personen (Hygienekontrollen), Mittweida 2–3 "
                "Personen (Analytik Boden/Wasser), Dozenten 5 (3 intern, 2 extern), "
                "Gütersloh mit Labor und Verwaltung.",
                "Gesamtzahl der Odoo-Nutzer noch zu erheben.",
            ],
            "golive": OFFEN,
            "budget": OFFEN,
        },

        "profil": [
            ["Bundesweites Hygieneinstitut, VDI-Schulungspartner (VDI 6022 / 2047)", "☒ ja",
             "Schulungen werden bundesweit beworben; Pflichtschulung für Großkonzerne und Industrieparks"],
            ["Inspektionen vor Ort: RLT-Anlagen, Verdunstungskühlanlagen, Luftentkeimung", "☒ ja",
             "Hygienekontrollen; 3er-Team am Standort München"],
            ["Beratung & Gefährdungsbeurteilungen (Trinkwasser-, Krankenhaushygiene)", "☐ offen",
             "im Termin nicht vertieft"],
            ["Bundesweite VDI-Schulungen, Termine online buchbar", "☒ ja",
             "Durchführung ab mehr als 10 Teilnehmern; Seminarfläche 6 Monate im Voraus gebucht"],
            ["Labor & Forschung, Entwicklung neuer Analyseverfahren", "☒ ja / teilweise",
             "Mikrobiologie mit wenigen Analysetechniken (siehe Laborprozess); Forschung nicht besprochen"],
            ["Standort Gütersloh (Elbrachtsweg 76, 33332)", "☒ ja",
             "Hauptstandort – ergänzt: Standorte München und Technologiepark Mittweida"],
            ["Ein Standort / eine Gesellschaft", "☐ nein",
             "Drei Standorte: Gütersloh, München, Technologiepark Mittweida. "
             "Gesellschaftsrechtliche Struktur zu klären"],
        ],

        "standorte": [
            ["Gütersloh", "Hauptstandort",
             "Labor, Bebrütung & Auszählung, Gutachtenerstellung, Verwaltung"],
            ["München", "2–3 Personen",
             "Hygienekontrollen; Büro in einem Privatgebäude"],
            ["Technologiepark Mittweida", "2–3 Personen",
             "Analytik von Boden und Wasser"],
        ],

        "apps": [
            ["CRM & Verkauf", "Anfragen, Angebote & Aufträge", "☒ ja",
             "Auftragseingang startet den Laborprozess"],
            ["Rechnungsstellung", "Rechnungen & Zahlungen", "☒ ja",
             "Rechnung nach Fertigstellung des Gutachtens"],
            ["Einkauf & Lager", "Labor- & Prüfmaterial", "☒ ja",
             "Gelatinefilter ca. 15 T€, Greiner-Röhrchen, 5–6 Händlerprodukte"],
            ["Außendienst", "Einsatzplanung & Prüfberichte vor Ort", "☒ ja",
             "Hygienekontrollen vor Ort; Fotos per Pixel 8"],
            ["Projekte & Zeiterfassung", "Beratungsaufträge & Gutachten", "☐ offen",
             "im Termin nicht vertieft"],
            ["Veranstaltungen & E-Learning", "VDI-Schulungen inkl. Online-Buchung", "☒ ja",
             "Mindestteilnehmerzahl 10; Vorplanung 6 Monate; 5 Dozenten; Inhouse-Schulungen"],
            ["Abonnements", "Wiederkehrende Prüfintervalle & Verträge", "☐ offen",
             "wiederkehrende Intervalle nicht besprochen"],
            ["Qualität & Dokumente", "Proben, Berichte & Nachweise", "☒ ja",
             "Proben, Platten, Fotos, Gutachten; Dokumentenlayout ist bindend"],
            ["Buchhaltung", "Finanzbuchhaltung, DATEV-Anbindung", "☐ offen",
             "heutiges Buchhaltungssystem nicht erhoben"],
            ["Produktion", "vorher nicht eingeplant", "☒ ja, minimal",
             "geringe Eigenproduktion einiger Chemikalien"],
            ["Website / E-Commerce", "vorher nicht eingeplant", "☐ offen",
             "Kurse werden bundesweit beworben – Buchungsweg zu klären"],
            ["Mehrgesellschaft / Multi-Company", "zwei Standorte abbilden", "☒ ja",
             "Gütersloh und München; Struktur und Konsolidierung zu klären"],
            ["Labor-/LIMS-Funktionen", "Ersatz für das Gutachtenprogramm", "☐ zu prüfen",
             "Auszählung, Fotodokumentation, Berichtserzeugung mit festem Layout"],
        ],

        "laborprozess": {
            "analytik": [
                "Proben kommen ins Labor und werden mikrobiologisch aufbereitet.",
                "Bestimmt wird die Keimzahl auf zwei Medien: Gesamtkeimzahl und Pilze.",
                "Keine große Bandbreite unterschiedlicher Analysetechniken.",
            ],
            "mengen": [
                "Bei voller Auslastung 100 Anlagen pro Woche.",
                "8–9 Platten je Anlage.",
                "Ergibt rund 800–900 Platten pro Woche. Im Termin wurde „9.000“ notiert – "
                "die Zahl ist zu verifizieren (Woche, Monat oder Jahr?).",
            ],
            "bebruetung": [
                "Proben liegen im Brutschrank und werden nach 3 oder 5 Tagen ausgezählt.",
                "Nach Sichtung durch das Labor werden die Proben autoklaviert und entsorgt.",
            ],
        },

        "schulungsgeschaeft": [
            "Schulungen werden bundesweit beworben.",
            "Eine Schulung findet statt, wenn mehr als 10 Teilnehmer zusammenkommen.",
            "Zielgruppe: Großkonzerne und Industrieparks – für diese ist es eine Pflichtschulung.",
            "VDI-Schulungen werden lange vorgeplant; die Seminarfläche wird 6 Monate im Voraus gebucht.",
            "Inhouse-Schulungen beim Kunden gehören ebenfalls zum Angebot.",
            "Dozenten: 3 intern und 2 extern, insgesamt 5. Ein zweites Team soll aufgebaut werden.",
        ],

        "warenwirtschaft": [
            "In der biotec wird nur minimal produziert – einige Chemikalien.",
            "Händlerprodukte: 5–6 Artikel.",
            "Lager: 2–3 Brutreaktoren (Begriff zu prüfen), Greiner-Röhrchen mit Chemikalien.",
            "Verbrauchsmaterial auf Lager: Gelatinefilter im Wert von rund 15.000 €.",
        ],

        "aufgaben": [
            ["1", "Protokoll & Zusammenfassung des Discovery Calls versenden", "CertoClav", "", "☐ offen"],
            ["2", "Rohdaten bereitstellen: Anlagenlisten, Kunden-/Artikelstamm, Kurskatalog",
             "biotec", "", "☐ offen"],
            ["3", "Gutachten- und Berichtsvorlagen als Muster bereitstellen (Layout ist bindend)",
             "biotec", "", "☐ offen"],
            ["4", "Delphi-Applikation: Exportmöglichkeiten und Schnittstellen dokumentieren",
             "biotec / CertoClav", "", "☐ offen"],
            ["5", "Anzahl und Namen der vor Ort zu schulenden Personen bestätigen (3 oder 4)",
             "biotec", "", "☐ offen"],
            ["6", "Hauptansprechpartner für die Iterationen benennen", "biotec", "", "☐ offen"],
            ["7", "Scoping & Angebot – inkl. Vor-Ort-Schulung statt Train-the-Trainer, "
                  "Multi-Company und Layouttreue", "CertoClav", "", "☐ offen"],
            ["8", "Klären, ob die Certania-Anforderungen (P2P/O2C, POC/WIP) zum Scope gehören",
             "CertoClav / Certania", "", "☐ offen"],
        ],

        "offene_punkte": [
            ["Anzahl der vor Ort zu schulenden Personen: 3 oder 4? Vierter Name fehlt", "biotec", ""],
            ["Plattenmenge: rund 900 oder 9.000 – und je Woche, Monat oder Jahr?", "biotec", ""],
            ["biotec Süd: welchem der drei Standorte entspricht der Name?", "biotec", ""],
            ["Nachnamen und Funktionen von „Stefan“ und „Markus“", "biotec", ""],
            ["Sind die drei Standorte eigene Gesellschaften oder Betriebsstätten? "
             "Konsolidierungsbedarf?", "biotec", ""],
            ["Welche Schnittstellen bedient die Delphi-Applikation heute?", "biotec", ""],
            ["Bleibt OneDrive die Dokumentenablage oder übernimmt Odoo Dokumente?", "gemeinsam", ""],
            ["Heutiges Buchhaltungssystem und DATEV-Anbindung", "biotec", ""],
            ["Michael Brand formell als Hauptansprechpartner bestätigen", "biotec", ""],
            ["Nicole Krupa: „Biotec Stetic“ – welche Einheit ist gemeint?", "biotec", ""],
            ["Westbomke: externer IT-Dienstleister oder eigene Abteilung? Rolle im Projekt "
             "(Zugänge, Schnittstellen, Delphi-Betreuung)?", "biotec", ""],
        ],

        "intern": [
            "Aufwandstreiber: (1) Layouttreue der Gutachten – Berichtsentwicklung in Odoo statt "
            "Standardvorlagen, (2) Ablösung der Delphi-Applikation samt unbekannter "
            "Schnittstellen, (3) Multi-Company für zwei Standorte, (4) Vor-Ort-Schulung von "
            "3–4 Personen zusätzlich zum Einrichtungsaufwand.",
            "Der Wegfall von Train-the-Trainer verschiebt Aufwand vom Kunden zu CertoClav – die "
            "Erzählung „Umstellung in einer Woche“ ist damit nicht haltbar und sollte im Angebot "
            "nicht wiederholt werden.",
            "Der Laborprozess ist der eigentliche Kern: Anlagenliste, Probenverfolgung, "
            "Auszählung mit Fotodokumentation, Berichtserzeugung. Vor dem Angebot prüfen, wie "
            "weit Odoo Qualität/Dokumente das ohne Eigenentwicklung tragen.",
            "Das Schulungsgeschäft mit Mindestteilnehmerzahl und 6 Monaten Vorlauf ist ein guter "
            "Quick-Win-Kandidat (Odoo Veranstaltungen) und liefert früh sichtbaren Nutzen.",
            "Entscheiderlage: Dr. Thomas Wilke ist CEO biotec und gleichzeitig Direktor Food bei "
            "Certania – die Freigabe läuft über ihn, nicht über Dr. Bermpohl (Prokura). Moritz "
            "Gruber bleibt Initiator auf Gesellschafterebene.",
            "Westbomke betreut die IT-Administration. Für die Ablösung der Delphi-Applikation und "
            "die Schnittstellen ist das die entscheidende technische Gegenstelle – früh "
            "einbinden, statt erst in der Umsetzungsphase.",
        ],
    },
}


# --------------------------------------------------------------------------- Dokument
def build(path, termin=None):
    d = TERMINE.get(termin) if termin else None
    antworten = (d or {}).get("antworten", {})
    doc = Document()
    seite_einrichten(doc, "Protokoll Discovery Call  ·  Dok.-Nr. DISC-2026-001")
    titel(doc, "Protokoll Discovery Call",
          "Odoo für die biotec GmbH  ·  Verstehen, wo Sie stehen – "
          "und prüfen, ob und wie Odoo Sie weiterbringt.")

    # Metadaten
    table(doc, ["Feld", "Eintrag"], [4.5, 12.5], d["meta"] if d else LEERE_META)

    # Teilnehmer
    if d:
        heading(doc, "Teilnehmer", kicker="Laut Einladung")
        table(doc, ["Name", "E-Mail", "Organisation", "Rolle / Funktion", "Status"],
              [3.3, 4.7, 2.8, 3.3, 2.9], d["teilnehmer"])
        if d.get("hinweis"):
            note(doc, d["hinweis"])

    # Kernaussagen
    heading(doc, "Kernaussagen des Gesprächs", kicker="Zusammenfassung")
    if d and d.get("kernaussagen"):
        bullets(doc, d["kernaussagen"])
    else:
        note(doc, "Drei bis fünf Sätze: Ausgangslage, wichtigste Erkenntnis, Einschätzung Machbarkeit.")
        answer_lines(doc, 5)

    # 1 CertoClav & Ansatz
    heading(doc, "1 · CertoClav & unser Ansatz", kicker="Agenda 01 – ca. 10 Min.")
    note(doc, "Vorgestellt: SPOC auf Beraterseite, Einrichtung + Datenimport durch CertoClav "
              "mit Claude, iterative Feedbackrunden, Train-the-Trainer.")
    question(doc, "Reaktionen, Rückfragen, Bedenken zum Arbeitsmodell",
             "(z. B. SPOC-Rolle, KI-Einsatz, Datenschutz, Verfügbarkeit)", lines=3,
             antwort=antworten.get("modell_reaktion"))

    # 2 Firmenprofil
    heading(doc, "2 · Firmenprofil-Check", kicker="Deck-Folie „Was wir über biotec wissen“")
    note(doc, "Unsere Annahmen aus biotec-gmbh.com – Korrekturen und Ergänzungen festhalten.")
    table(doc, ["Annahme", "Bestätigt?", "Korrektur / Ergänzung"], [6.0, 2.6, 8.4],
          d["profil"] if d and d.get("profil") else PROFIL_LEER)
    if d and d.get("standorte"):
        heading(doc, "2b · Standorte & Struktur", kicker="Ergänzung aus dem Termin")
        table(doc, ["Standort", "Team", "Aufgaben"], [4.6, 3.0, 9.4], d["standorte"])
    else:
        question(doc, "Was haben wir übersehen? Was ist am wichtigsten – was läuft heute schon gut?",
                 lines=3)

    # 3 Discovery 1/3
    heading(doc, "3 · Schmerzpunkte & Ziele", kicker="Discovery 1/3")
    for key, q, hint in F_SCHMERZ:
        question(doc, q, hint, antwort=antworten.get(key))

    # 4 Discovery 2/3
    heading(doc, "4 · Prozesse, Systeme & Daten", kicker="Discovery 2/3")
    note(doc, "Basis für Scoping, Phasierung und den Claude-gestützten Datenimport.")
    for key, q, hint in F_PROZESSE:
        question(doc, q, hint, antwort=antworten.get(key))

    # 4b Laborprozess, Schulungen, Warenwirtschaft
    if d and d.get("laborprozess"):
        lp = d["laborprozess"]
        heading(doc, "4b · Laborprozess & Mengen", kicker="Ergänzung aus dem Termin")
        for block_titel, key in [("Analytik", "analytik"), ("Bebrütung & Entsorgung", "bebruetung"),
                           ("Mengengerüst", "mengen")]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(0)
            p.add_run(block_titel).font.bold = True
            bullets(doc, lp[key])

    if d and d.get("schulungsgeschaeft"):
        heading(doc, "4c · Schulungsgeschäft", kicker="Ergänzung aus dem Termin")
        bullets(doc, d["schulungsgeschaeft"])

    if d and d.get("warenwirtschaft"):
        heading(doc, "4d · Produktion, Handelsware & Lager", kicker="Ergänzung aus dem Termin")
        bullets(doc, d["warenwirtschaft"])

    doc.add_page_break()

    # 5 App-Hypothese
    heading(doc, "5 · App-Hypothese validieren", kicker="Deck-Folie „Odoo-Apps für biotec“")
    table(doc, ["Odoo-App", "Zweck (Hypothese)", "Relevant?", "Anmerkung"],
          [3.6, 5.4, 2.6, 5.4], d["apps"] if d and d.get("apps") else APPS_LEER)

    # 6 Certania-Anforderungen
    heading(doc, "6 · Anforderungen Finance / Certania-Gruppe", kicker="Ergänzung zum Deck")
    note(doc, "Aus dem Gruppen-Strang (Jeannette Bühler, Head of Group Accounting): "
              "vollintegriertes P2P/O2C und POC-/WIP-Bewertung. Klären, ob das zum Scope gehört.")
    for key, q, hint in F_FINANCE:
        question(doc, q, hint, antwort=antworten.get(key))

    # 7 Discovery 3/3
    heading(doc, "7 · Team, Zeitrahmen & Budget", kicker="Discovery 3/3")
    for key, q, hint in F_TEAM:
        question(doc, q, hint, antwort=antworten.get(key))

    # 8 Nächste Schritte
    heading(doc, "8 · Nächste Schritte & Aufgaben", kicker="Deck-Folie „Nächste Schritte“")
    table(doc, ["Nr.", "Aufgabe", "Wer", "Termin", "Status"], [1.2, 8.0, 3.0, 2.4, 2.4],
          d["aufgaben"] if d and d.get("aufgaben") else AUFGABEN_LEER)

    # 9 Offene Punkte
    heading(doc, "9 · Offene Punkte & Risiken")
    note(doc, "Was im Termin nicht geklärt werden konnte – inkl. wer es klärt.")
    table(doc, ["Offener Punkt", "Klärt", "Bis"], [11.0, 3.0, 3.0],
          d["offene_punkte"] if d and d.get("offene_punkte") else [["", "", ""] for _ in range(6)])

    # 10 Interne Notizen
    heading(doc, "10 · Interne Notizen", kicker="Nicht Teil des Kundenprotokolls")
    note(doc, "Einschätzung Machbarkeit, Aufwandstreiber, Stakeholder-Dynamik, Preisindikation. "
              "Vor dem Versand an den Kunden entfernen.")
    if d and d.get("intern"):
        bullets(doc, d["intern"])
    else:
        answer_lines(doc, 6)

    doc.save(path)
    print(f"geschrieben: {path}")


if __name__ == "__main__":
    ziel = sys.argv[1] if len(sys.argv) > 1 else "Protokoll_Discovery_Call.docx"
    termin = sys.argv[2] if len(sys.argv) > 2 else None
    if termin and termin not in TERMINE:
        sys.exit(f"Unbekannter Termin '{termin}'. Verfügbar: {', '.join(TERMINE)}")
    build(ziel, termin)
