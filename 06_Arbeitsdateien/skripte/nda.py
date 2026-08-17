#!/usr/bin/env python3
"""Erzeugt die Vertraulichkeitsvereinbarung (NDA) CertoClav ↔ biotec.

Gegenseitige Vereinbarung: biotec offenbart Kundendaten, Gutachten, kaufmännische
Zahlen sowie Datenbank und Quellcode des Altsystems; CertoClav offenbart Konzepte,
Kalkulationsgrundlagen und Vorgehensweisen.

WICHTIG: Vorlage, keine Rechtsberatung. Vor Verwendung juristisch prüfen lassen –
insbesondere Rechtswahl, Gerichtsstand und Laufzeit.

Aufruf:  python3 nda.py <zieldatei.docx>
"""
import sys

from docx import Document
from docx.shared import Cm, Pt

from docx_bausteine import ACCENT, GREY, bottom_border, heading, note, seite_einrichten, table, titel

# Rechtswahl – Vorbelegung: deutsches Recht, weil biotec die Hauptoffenbarende Partei
# und in Deutschland ansässig ist. Alternative: österreichisches Recht, Gerichtsstand Linz.
RECHT = "Recht der Bundesrepublik Deutschland"
GERICHTSSTAND = "Gütersloh"
LAUFZEIT_JAHRE = 3


def para(doc, text, groesse=10, abstand=6, bold=False, kursiv=False, einzug=0.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(abstand)
    if einzug:
        p.paragraph_format.left_indent = Cm(einzug)
    r = p.add_run(text)
    r.font.size = Pt(groesse)
    r.font.bold = bold
    r.font.italic = kursiv
    return p


def ziffer(doc, nummer, text):
    """Nummerierter Unterabsatz, z. B. 3.1."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(1.1)
    p.paragraph_format.first_line_indent = Cm(-1.1)
    r = p.add_run(nummer + "\t")
    r.font.size = Pt(10)
    r.font.bold = True
    r2 = p.add_run(text)
    r2.font.size = Pt(10)


def paragraf(doc, nummer, titel_text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    r = h.add_run(f"{nummer}. {titel_text}")
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = ACCENT


def build(path):
    doc = Document()
    seite_einrichten(doc, "Vertraulichkeitsvereinbarung  ·  Dok.-Nr. NDA-2026-001")
    titel(doc, "Vertraulichkeitsvereinbarung",
          "Gegenseitige Vereinbarung über die Behandlung vertraulicher Informationen")

    para(doc, "zwischen", groesse=10, abstand=4)
    table(doc, ["Partei", "Angaben"], [3.2, 13.8], [
        ["biotec GmbH\n(nachfolgend „biotec“)",
         "Elbrachtsweg 76, 33332 Gütersloh, Deutschland\n"
         "Telefon +49 5241 307200 · info@biotec-gmbh.com\n"
         "vertreten durch: ______________________________"],
        ["CertoClav Sterilizer GmbH\n(nachfolgend „CertoClav“)",
         "Peintner Straße 10, 4060 Leonding, Österreich\n"
         "Telefon +43 732 674 278 · support@certoclav.com\n"
         "Firmenbuch: Landesgericht Linz, FN 122912d · UID ATU22821702\n"
         "vertreten durch: Michael Simon (geb. Dirix), MSc., Geschäftsführer"],
    ])
    para(doc, "– gemeinsam die „Parteien“, jeweils einzeln „Partei“ –", groesse=9.5,
         kursiv=True, abstand=10)

    # 1 Zweck
    paragraf(doc, 1, "Zweck")
    ziffer(doc, "1.1", "Die Parteien prüfen und führen die Einführung des ERP-Systems Odoo "
                       "bei biotec durch (der „Zweck“). Dazu ist der Austausch vertraulicher "
                       "Informationen erforderlich.")
    ziffer(doc, "1.2", "Vertrauliche Informationen dürfen ausschließlich für diesen Zweck "
                       "verwendet werden.")

    # 2 Vertrauliche Informationen
    paragraf(doc, 2, "Vertrauliche Informationen")
    ziffer(doc, "2.1", "Vertrauliche Informationen sind alle Informationen, die eine Partei "
                       "(„offenbarende Partei“) der anderen („empfangende Partei“) im "
                       "Zusammenhang mit dem Zweck zugänglich macht – unabhängig von Form "
                       "und Datenträger, mündlich, schriftlich, elektronisch oder durch "
                       "Einsichtnahme.")
    ziffer(doc, "2.2", "Dazu gehören insbesondere, ohne darauf beschränkt zu sein:")
    for t in ("Kunden-, Lieferanten- und Kontaktdaten sowie Anlagen- und Objektdaten,",
              "Prüfberichte, Gutachten, Analyseergebnisse, Prüfpläne und Bewertungsschemata,",
              "kaufmännische Daten: Preise, Kalkulationen, Umsätze, Konten, offene Posten,",
              "Datenbanken, Datenbankschemata, Datenexporte und Quellcode einschließlich "
              "Dokumentation der von biotec eingesetzten Software,",
              "Vorlagen, Layouts und Textbausteine der Dokumente von biotec,",
              "Konzepte, Architekturen, Vorgehensmodelle, Aufwandsschätzungen und "
              "Kalkulationsgrundlagen von CertoClav,",
              "Verträge, Vertragsentwürfe und Konditionen beider Parteien."):
        para(doc, "– " + t, groesse=10, abstand=2, einzug=1.6)
    ziffer(doc, "2.3", "Eine Kennzeichnung als „vertraulich“ ist nicht erforderlich. Die "
                       "Parteien behandeln die genannten Informationen auch ohne Kennzeichnung "
                       "als vertraulich.")
    ziffer(doc, "2.4", "Die Parteien sind sich einig, dass die in dieser Vereinbarung "
                       "getroffenen Maßnahmen angemessene Geheimhaltungsmaßnahmen im Sinne "
                       "des Geschäftsgeheimnisgesetzes darstellen.")

    # 3 Ausnahmen
    paragraf(doc, 3, "Ausnahmen")
    ziffer(doc, "3.1", "Nicht vertraulich sind Informationen, die")
    for t in ("bei Erhalt bereits allgemein bekannt waren oder ohne Verletzung dieser "
              "Vereinbarung allgemein bekannt werden,",
              "der empfangenden Partei vor der Offenbarung rechtmäßig und ohne "
              "Vertraulichkeitsbindung bekannt waren,",
              "die empfangende Partei ohne Nutzung der vertraulichen Informationen selbst "
              "entwickelt hat,",
              "von einem berechtigten Dritten ohne Vertraulichkeitsbindung überlassen wurden."):
        para(doc, "– " + t, groesse=10, abstand=2, einzug=1.6)
    ziffer(doc, "3.2", "Besteht eine gesetzliche oder behördliche Pflicht zur Offenlegung, "
                       "informiert die empfangende Partei die offenbarende Partei "
                       "unverzüglich und beschränkt die Offenlegung auf das erforderliche Maß.")

    # 4 Pflichten
    paragraf(doc, 4, "Pflichten der empfangenden Partei")
    ziffer(doc, "4.1", "Die empfangende Partei bewahrt vertrauliche Informationen geheim, "
                       "schützt sie mit der Sorgfalt, die sie eigenen vertraulichen "
                       "Informationen entgegenbringt – mindestens jedoch mit "
                       "verkehrsüblicher Sorgfalt – und gibt sie nicht an Dritte weiter.")
    ziffer(doc, "4.2", "Zugang erhalten nur Personen, die die Informationen für den Zweck "
                       "benötigen (Kenntnis-nur-wenn-nötig) und die zur Vertraulichkeit "
                       "verpflichtet sind. Dies umfasst Mitarbeitende, Organe und "
                       "Beauftragte.")
    ziffer(doc, "4.3", "Die Einbeziehung von Unterauftragnehmern ist zulässig, wenn diese "
                       "schriftlich zu einer mindestens gleichwertigen Vertraulichkeit "
                       "verpflichtet werden. Die empfangende Partei bleibt verantwortlich.")
    ziffer(doc, "4.4", "Vertrauliche Informationen werden nicht länger und nicht weiter "
                       "gespeichert als für den Zweck erforderlich. Die Speicherung erfolgt "
                       "auf zugriffsbeschränkten Systemen.")
    ziffer(doc, "4.5", "Quellcode und Datenbankinhalte werden ausschließlich zur Analyse und "
                       "zur Datenübernahme im Rahmen des Zwecks verwendet. Eine Nutzung für "
                       "eigene Produkte oder für Dritte ist ausgeschlossen.")

    doc.add_page_break()

    # 5 Datenschutz
    paragraf(doc, 5, "Personenbezogene Daten")
    ziffer(doc, "5.1", "Soweit vertrauliche Informationen personenbezogene Daten enthalten, "
                       "beschränken die Parteien deren Übermittlung auf das für den Zweck "
                       "erforderliche Maß. biotec stellt Beispieldaten soweit möglich "
                       "anonymisiert oder pseudonymisiert bereit.")
    ziffer(doc, "5.2", "Diese Vereinbarung ersetzt keinen Vertrag über die "
                       "Auftragsverarbeitung. Sofern CertoClav personenbezogene Daten im "
                       "Auftrag von biotec verarbeitet, schließen die Parteien zusätzlich "
                       "eine Vereinbarung nach Art. 28 DSGVO.")

    # 6 Kein Rechteübergang
    paragraf(doc, 6, "Keine Rechteübertragung, keine Abschlusspflicht")
    ziffer(doc, "6.1", "Vertrauliche Informationen bleiben Eigentum der offenbarenden Partei. "
                       "Diese Vereinbarung begründet keine Lizenz und keine sonstigen Rechte "
                       "an gewerblichen Schutzrechten, Urheberrechten oder Know-how.")
    ziffer(doc, "6.2", "Aus dieser Vereinbarung ergibt sich keine Verpflichtung, einen "
                       "weiteren Vertrag abzuschließen oder Informationen offenzulegen.")
    ziffer(doc, "6.3", "Die offenbarende Partei übernimmt keine Gewähr für Richtigkeit oder "
                       "Vollständigkeit der offenbarten Informationen.")

    # 7 Rückgabe
    paragraf(doc, 7, "Rückgabe und Löschung")
    ziffer(doc, "7.1", "Auf Verlangen der offenbarenden Partei, spätestens nach Erledigung "
                       "des Zwecks, gibt die empfangende Partei alle vertraulichen "
                       "Informationen samt Kopien zurück oder löscht sie und bestätigt dies "
                       "auf Wunsch schriftlich.")
    ziffer(doc, "7.2", "Ausgenommen sind Informationen, für die eine gesetzliche "
                       "Aufbewahrungspflicht besteht, sowie Sicherungskopien aus "
                       "automatisierten Backup-Verfahren. Für diese gilt die "
                       "Vertraulichkeitspflicht fort.")
    ziffer(doc, "7.3", "Datenbanksicherungen und Quellcode werden nach Abschluss des "
                       "Projekts gelöscht, sofern die Parteien nichts anderes vereinbaren.")

    # 8 Laufzeit
    paragraf(doc, 8, "Laufzeit")
    ziffer(doc, "8.1", "Diese Vereinbarung tritt mit Unterzeichnung durch beide Parteien in "
                       "Kraft.")
    ziffer(doc, "8.2", f"Die Vertraulichkeitspflichten gelten für {LAUFZEIT_JAHRE} Jahre "
                       "nach Beendigung der Zusammenarbeit fort. Für Informationen, die "
                       "Geschäftsgeheimnisse darstellen, gelten sie unbefristet, solange "
                       "die Voraussetzungen eines Geschäftsgeheimnisses vorliegen.")

    # 9 Schluss
    paragraf(doc, 9, "Schlussbestimmungen")
    ziffer(doc, "9.1", "Änderungen und Ergänzungen bedürfen der Schriftform. Das gilt auch "
                       "für die Änderung dieser Klausel.")
    ziffer(doc, "9.2", "Sollte eine Bestimmung unwirksam sein, bleibt die Vereinbarung im "
                       "Übrigen wirksam. Die Parteien ersetzen die unwirksame Bestimmung "
                       "durch eine wirksame, die dem verfolgten Zweck am nächsten kommt.")
    ziffer(doc, "9.3", f"Es gilt das {RECHT} unter Ausschluss des UN-Kaufrechts. "
                       f"Ausschließlicher Gerichtsstand ist {GERICHTSSTAND}.")
    ziffer(doc, "9.4", "Diese Vereinbarung enthält die vollständige Abrede der Parteien zur "
                       "Vertraulichkeit und ersetzt frühere Absprachen hierzu.")

    # Unterschriften
    heading(doc, "Unterschriften")
    t = table(doc, ["biotec GmbH", "CertoClav Sterilizer GmbH"], [8.5, 8.5], [
        ["\n\nOrt, Datum\n\n\n\n\n________________________________\nName, Funktion",
         "\n\nOrt, Datum\n\n\n\n\n________________________________\n"
         "Michael Simon (geb. Dirix), MSc.\nGeschäftsführer"],
    ])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    bottom_border(p)
    r = p.add_run("Hinweis für den internen Gebrauch")
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = GREY
    note(doc, "Diese Vereinbarung ist eine Vorlage und keine Rechtsberatung. Vor Verwendung "
              "juristisch prüfen lassen. Vorbelegt sind deutsches Recht und Gerichtsstand "
              "Gütersloh, weil biotec die hauptsächlich offenbarende Partei ist – das "
              "erleichtert die Zeichnung auf Kundenseite. Alternative: österreichisches "
              "Recht, Gerichtsstand Linz. Diesen Hinweis vor dem Versand löschen.")

    doc.save(path)
    print(f"geschrieben: {path}")


if __name__ == "__main__":
    build(sys.argv[1] if len(sys.argv) > 1 else "NDA_CertoClav_biotec.docx")
